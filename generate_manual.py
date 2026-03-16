#!/usr/bin/env python3
"""Generate Apotheca Marketing Suite User Manual Part 1 as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style Configuration ──────────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# List style
list_style = doc.styles['List Number']
list_style.font.name = 'Calibri'
list_style.font.size = Pt(11)


def add_para(text, bold=False, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_callout(prefix, text):
    """Add a TIP: or IMPORTANT: callout as a bold-prefixed paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{prefix}: ')
    r1.bold = True
    if prefix == 'IMPORTANT':
        r1.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    else:
        r1.font.color.rgb = RGBColor(0x00, 0x70, 0x30)
    p.add_run(text)
    return p


def add_steps(steps):
    """Add a numbered step list."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'Step {i}. ')
        r.bold = True
        p.add_run(step)


def add_bullet_list(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_table_row_data(table, rows):
    """Populate table rows. First row is header."""
    for i, row_data in enumerate(rows):
        if i == 0:
            row = table.rows[0]
        else:
            row = table.add_row()
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Shade header cells
                shading = cell._element.get_or_add_tcPr()
                sh = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): '1B3A5C',
                    qn('w:val'): 'clear',
                })
                shading.append(sh)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Apotheca® Marketing Suite')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('User Manual — Part 1')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run('Version 1.0.0')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

chapters_note = doc.add_paragraph()
chapters_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = chapters_note.add_run('Chapters 1–4: Installation, Settings, GDPR & Compliance, Subscriber Management')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Chapter 1: Installation & Activation',
    '    1.1 System Requirements',
    '    1.2 Uploading and Installing the Plugin',
    '    1.3 Activating the Plugin',
    '    1.4 What Happens on Activation',
    '    1.5 Verifying Successful Installation',
    '    1.6 Admin Menu Overview',
    '',
    'Chapter 2: Global Settings & Configuration',
    '    2.1 Navigating to the Settings Page',
    '    2.2 Checkout Opt-In Settings',
    '    2.3 Registration Subscriber Capture',
    '    2.4 GDPR Double Opt-In',
    '    2.5 Consent Text',
    '    2.6 Abandoned Cart Timeout',
    '    2.7 Frequency Caps',
    '    2.8 Send Window',
    '    2.9 Attribution Window',
    '    2.10 Unsubscribe Page Customisation',
    '    2.11 Uninstall Data Option',
    '    2.12 Saving Your Settings',
    '',
    'Chapter 3: GDPR & Compliance',
    '    3.1 Understanding Double Opt-In',
    '    3.2 The Confirmation Email Flow',
    '    3.3 The Unsubscribe Endpoint',
    '    3.4 How Tokenised Unsubscribe Links Work',
    '    3.5 Consent Timestamp Tracking',
    '    3.6 One-Click Unsubscribe for Subscribers',
    '',
    'Chapter 4: Subscriber Management',
    '    4.1 How Subscribers Are Captured',
    '    4.2 Viewing the Subscriber List',
    '    4.3 Searching and Filtering Subscribers',
    '    4.4 Understanding the Subscriber Table Columns',
    '    4.5 Subscriber Statuses Explained',
    '    4.6 Subscriber Sources',
    '    4.7 Tags',
    '    4.8 RFM Score and Segment',
    '    4.9 Subscriber Profile Page',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('    '):
        p = doc.add_paragraph(item.strip())
        p.paragraph_format.left_indent = Inches(0.4)
    else:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: INSTALLATION & ACTIVATION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Chapter 1: Installation & Activation', level=1)

add_para(
    'This chapter walks you through installing and activating the Apotheca® Marketing Suite plugin on your WordPress site. '
    'By the end, you will have the plugin running and ready for configuration.'
)

# 1.1
doc.add_heading('1.1 System Requirements', level=2)

add_para('Before you install Apotheca® Marketing Suite, make sure your hosting environment meets the following minimum requirements:')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row_data(table, [
    ['Requirement', 'Minimum Version'],
    ['WordPress', '6.4 or higher'],
    ['PHP', '8.0 or higher'],
    ['WooCommerce', '8.0 – 9.0 (must be installed and active)'],
    ['MySQL / MariaDB', '5.7 / 10.3 or higher (standard WordPress requirement)'],
    ['Web Server', 'Apache or Nginx (standard WordPress hosting)'],
])

add_para('')
add_callout('IMPORTANT', 'WooCommerce must be installed and activated before you activate Apotheca® Marketing Suite. '
            'If WooCommerce is not active, the plugin will refuse to activate and display an error message.')

add_callout('TIP', 'You can check your PHP version by going to Tools → Site Health → Info → Server in your WordPress admin.')

# 1.2
doc.add_heading('1.2 Uploading and Installing the Plugin', level=2)

add_para('You can install the plugin by uploading the ZIP file through the WordPress admin panel.')

add_steps([
    'Log in to your WordPress admin dashboard.',
    'In the left sidebar, navigate to Plugins → Add New.',
    'Click the Upload Plugin button at the top of the page.',
    'Click Choose File and select the apotheca-marketing-suite.zip file from your computer.',
    'Click Install Now.',
    'WordPress will upload and extract the plugin files. Wait for the "Plugin installed successfully" message to appear.',
])

add_callout('TIP', 'If you prefer FTP/SFTP installation, extract the ZIP file and upload the entire '
            'apotheca-marketing-suite folder to your /wp-content/plugins/ directory.')

# 1.3
doc.add_heading('1.3 Activating the Plugin', level=2)

add_steps([
    'After installation, click the Activate Plugin link on the success page. Alternatively, go to Plugins → Installed Plugins and find "Apotheca® Marketing Suite" in the list.',
    'Click Activate.',
    'The plugin will run its activation routine automatically (see Section 1.4 for details).',
    'Once activated, you will see a new menu item called Apotheca® Marketing in the WordPress admin sidebar.',
])

add_callout('IMPORTANT', 'If WooCommerce is not active when you try to activate the plugin, you will see an error page that reads: '
            '"Apotheca® Marketing Suite requires WooCommerce to be installed and active." '
            'Use the Back button, install/activate WooCommerce first, and then try again.')

# 1.4
doc.add_heading('1.4 What Happens on Activation', level=2)

add_para('When you activate the plugin for the first time, the following actions occur automatically:')

doc.add_heading('1.4.1 Database Tables Created', level=3)

add_para(
    'The plugin creates 14 custom database tables in your WordPress database. '
    'These tables store all subscriber data, marketing events, automation flows, campaigns, and analytics. '
    'Each table name uses the prefix ams_ (appended to your WordPress table prefix).'
)

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
add_table_row_data(table2, [
    ['Table Name', 'Purpose'],
    ['ams_subscribers', 'Stores all subscriber records (email, name, phone, status, RFM data, consent info)'],
    ['ams_events', 'Tracks subscriber activity events (orders, page views, clicks)'],
    ['ams_flows', 'Stores automation flow definitions and their trigger configurations'],
    ['ams_flow_steps', 'Individual steps within each automation flow (emails, delays, conditions)'],
    ['ams_flow_enrolments', 'Tracks which subscribers are enrolled in which flows and their progress'],
    ['ams_campaigns', 'One-time email and SMS campaign records'],
    ['ams_segments', 'Subscriber segment definitions and conditions'],
    ['ams_sends', 'Log of every email and SMS sent, with open/click/bounce tracking'],
    ['ams_forms', 'Pop-up form configurations (design, targeting, trigger rules)'],
    ['ams_attributions', 'Revenue attribution records linking sends to WooCommerce orders'],
    ['ams_analytics_daily', 'Aggregated daily metrics for the analytics dashboard'],
    ['ams_ai_log', 'Log of AI feature usage (subject line generation, send-time optimisation)'],
    ['ams_reviews_cache', 'Cached product review data for review request flows'],
    ['ams_sync_inbound_log', 'Log of inbound data sync events from connected sites'],
])

doc.add_heading('1.4.2 Default Settings Initialised', level=3)

add_para(
    'The plugin saves a full set of default settings to your WordPress options table under the key ams_settings. '
    'You can customise all of these later from the Settings page (see Chapter 2).'
)

doc.add_heading('1.4.3 Rewrite Rules Flushed', level=3)

add_para(
    'WordPress rewrite rules are flushed to register the plugin\'s public endpoint for unsubscribe handling '
    '(see Chapter 3 for details on the /ams-unsubscribe/ endpoint).'
)

doc.add_heading('1.4.4 Database Version Recorded', level=3)

add_para(
    'The plugin stores its current database schema version (1.6.0) in the WordPress options table. '
    'On future plugin updates, the plugin compares this value with the bundled version and automatically runs any necessary database migrations.'
)

# 1.5
doc.add_heading('1.5 Verifying Successful Installation', level=2)

add_para('After activation, verify that everything installed correctly:')

add_steps([
    'Look at the WordPress admin sidebar. You should see a new menu item called Apotheca® Marketing with an email icon.',
    'Click Apotheca® Marketing → Dashboard. The dashboard page should load without errors.',
    'Click Apotheca® Marketing → Settings. Confirm that the settings page displays a form with fields for Checkout Opt-In, Double Opt-In, Frequency Caps, and other options.',
    'Click Apotheca® Marketing → Subscribers. You should see an empty subscriber table with column headers: Email, Name, Status, Source, Orders, Total Spent, RFM, and Subscribed.',
])

add_callout('TIP', 'If you see a white screen or a PHP error after activation, check that your server meets the minimum '
            'requirements listed in Section 1.1. The most common issue is running a PHP version older than 8.0.')

# 1.6
doc.add_heading('1.6 Admin Menu Overview', level=2)

add_para(
    'The Apotheca® Marketing top-level menu appears in the WordPress admin sidebar at position 56 '
    '(just below the WooCommerce menu). It uses the dashicons-email-alt icon. '
    'Clicking it reveals the following submenu pages:'
)

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
add_table_row_data(table3, [
    ['Submenu Item', 'Description'],
    ['Dashboard', 'Overview of your marketing performance and key metrics'],
    ['Subscribers', 'View, search, and manage your subscriber list'],
    ['Flows', 'Create and manage automated email/SMS flows'],
    ['Campaigns', 'Create and send one-time email and SMS campaigns'],
    ['Segments', 'Define subscriber segments based on conditions and behaviours'],
    ['Forms', 'Design and manage pop-up opt-in forms'],
    ['SMS', 'Configure SMS messaging settings and view SMS history'],
    ['Analytics', 'Detailed analytics dashboard with revenue attribution'],
    ['AI Settings', 'Configure AI-powered features (subject lines, send-time optimisation)'],
    ['Email Editor', 'Visual email template editor'],
    ['Reviews', 'Product review collection and review gate settings'],
    ['Sync', 'Multi-site synchronisation settings'],
    ['Settings', 'Global plugin settings (covered in Chapter 2)'],
])

add_callout('TIP', 'You need the manage_woocommerce capability to access any Apotheca® Marketing menu page. '
            'By default, this means you must be a Shop Manager or Administrator.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: GLOBAL SETTINGS & CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Chapter 2: Global Settings & Configuration', level=1)

add_para(
    'The Settings page is where you configure the core behaviour of Apotheca® Marketing Suite. '
    'Every setting on this page has a sensible default, but you should review each one to match your store\'s needs and compliance requirements.'
)

# 2.1
doc.add_heading('2.1 Navigating to the Settings Page', level=2)

add_steps([
    'Log in to your WordPress admin dashboard.',
    'In the left sidebar, click Apotheca® Marketing to expand the submenu.',
    'Click Settings.',
    'The settings page loads with a form containing all global configuration options, organised in a table layout.',
])

# 2.2
doc.add_heading('2.2 Checkout Opt-In Settings', level=2)

doc.add_heading('2.2.1 Checkout Opt-In Toggle', level=3)

add_para('Field name: Checkout Opt-In')
add_para('Type: Checkbox')
add_para('Default: Enabled (checked)')

add_para(
    'When enabled, a marketing consent checkbox is displayed on the WooCommerce checkout page, '
    'just above the "Place Order" button. Customers who check this box will be added to your subscriber list '
    'when they complete their order.'
)

add_para(
    'When disabled, no opt-in checkbox appears at checkout. Customers who place orders will still be tracked '
    'in the subscribers table (with a status of "never_subscribed"), but they will not receive marketing emails.'
)

add_callout('TIP', 'Even if you disable the checkout opt-in, you can still capture subscribers through registration (see Section 2.3) '
            'and pop-up forms (covered in a later chapter).')

doc.add_heading('2.2.2 Checkout Label Text', level=3)

add_para('Field name: Checkout Label')
add_para('Type: Text input (large)')
add_para('Default: "Keep me updated with news and offers via email."')

add_para(
    'This is the text that appears next to the opt-in checkbox on the checkout page. '
    'You can customise this to match your brand voice. The text should clearly communicate that '
    'the customer is opting in to receive marketing communications.'
)

add_callout('IMPORTANT', 'If you operate in the EU or other GDPR-regulated jurisdictions, '
            'make sure your checkout label text is clear and unambiguous about what the customer is consenting to. '
            'Avoid pre-checking the box by regulation — note that the plugin does pre-check the box by default for convenience, '
            'so if you are subject to GDPR, consider adjusting this behaviour or relying on double opt-in (Section 2.4).')

# 2.3
doc.add_heading('2.3 Registration Subscriber Capture', level=2)

add_para('Field name: Registration Capture')
add_para('Type: Checkbox')
add_para('Default: Enabled (checked)')

add_para(
    'When enabled, any user who registers a new account through WooCommerce is automatically added to your subscriber list. '
    'The subscriber record is created using the email address, first name, and last name from the WordPress user profile.'
)

add_para(
    'The subscriber\'s source is recorded as "registration" so you can identify how each subscriber was captured.'
)

add_para(
    'If double opt-in is enabled (see Section 2.4), newly registered users will receive a confirmation email '
    'and will not be marked as active subscribers until they click the confirmation link.'
)

add_callout('TIP', 'This feature works with WooCommerce\'s built-in "My Account" registration form as well as '
            'any registration triggered via the user_register WordPress hook.')

# 2.4
doc.add_heading('2.4 GDPR Double Opt-In', level=2)

add_para('Field name: Double Opt-In')
add_para('Type: Checkbox')
add_para('Default: Disabled (unchecked)')

add_para(
    'When enabled, new subscribers are not immediately marked as "subscribed." Instead, they are saved with a status '
    'of "pending" and receive a confirmation email. The subscriber must click the confirmation link in that email '
    'to activate their subscription. This is known as double opt-in and is the recommended approach for GDPR compliance.'
)

add_para('Here is how the flow works:')

add_steps([
    'A customer opts in at checkout, registers an account, or submits a form.',
    'The plugin creates a subscriber record with status "pending".',
    'A confirmation email is sent to the subscriber with a unique tokenised link.',
    'The subscriber clicks the confirmation link.',
    'The plugin updates the subscriber\'s status from "pending" to "subscribed", records the GDPR consent timestamp, and sets the gdpr_consent flag to true.',
    'A confirmation page is displayed: "Your subscription has been confirmed. Thank you!"',
])

add_callout('IMPORTANT', 'If you sell to customers in the EU, enabling double opt-in is strongly recommended. '
            'It provides verifiable proof that the subscriber actively confirmed their consent.')

add_callout('TIP', 'If you operate exclusively in regions where double opt-in is not required (e.g., the United States), '
            'you can leave this disabled for a simpler subscriber experience.')

# 2.5
doc.add_heading('2.5 Consent Text', level=2)

add_para('Field name: Consent Text')
add_para('Type: Textarea')
add_para('Default: "I agree to receive marketing emails. You can unsubscribe at any time."')

add_para(
    'This text is stored as part of your GDPR configuration. It represents the consent statement that is '
    'presented to subscribers when they opt in. You should customise this to match your privacy policy and comply '
    'with local regulations.'
)

# 2.6
doc.add_heading('2.6 Abandoned Cart Timeout', level=2)

add_para('Field name: Abandoned Cart Timeout (min)')
add_para('Type: Number input')
add_para('Default: 60 minutes')
add_para('Minimum: 15 minutes')

add_para(
    'This setting controls how long the plugin waits before considering a shopping cart "abandoned." '
    'If a customer adds items to their cart and does not complete checkout within this number of minutes, '
    'the cart is flagged as abandoned and can trigger an abandoned cart automation flow.'
)

add_para(
    'For example, with the default setting of 60 minutes: if a customer adds products to their cart at 2:00 PM '
    'and has not checked out by 3:00 PM, their cart is considered abandoned.'
)

add_callout('TIP', 'A value between 30 and 90 minutes works well for most stores. '
            'Setting it too short may flag carts that the customer was still intending to complete. '
            'Setting it too long delays your recovery emails.')

# 2.7
doc.add_heading('2.7 Frequency Caps', level=2)

add_para(
    'Frequency caps prevent your subscribers from receiving too many messages in a single day. '
    'This protects your sender reputation and avoids annoying your customers.'
)

doc.add_heading('2.7.1 Email Frequency Cap', level=3)

add_para('Field name: Email Frequency Cap (per 24h)')
add_para('Type: Number input')
add_para('Default: 3 emails per day')
add_para('Minimum: 1')

add_para(
    'This is the maximum number of marketing emails a single subscriber can receive within a 24-hour period. '
    'If a subscriber has already received 3 emails today (with the default setting), any additional emails '
    'scheduled for that subscriber will be held until the next day.'
)

doc.add_heading('2.7.2 SMS Frequency Cap', level=3)

add_para('Field name: SMS Frequency Cap (per 24h)')
add_para('Type: Number input')
add_para('Default: 2 SMS messages per day')
add_para('Minimum: 1')

add_para(
    'This is the maximum number of marketing SMS messages a single subscriber can receive within a 24-hour period. '
    'SMS messages are more intrusive than emails, so the default cap is lower.'
)

add_callout('IMPORTANT', 'Frequency caps apply across all channels independently. A subscriber could receive '
            'up to 3 emails AND 2 SMS messages in a single day with the default settings. '
            'This includes messages from both campaigns and automation flows.')

add_callout('TIP', 'If you are running a major sale and want to send more emails than usual, '
            'you can temporarily increase the email frequency cap — but remember to lower it again afterward.')

# 2.8
doc.add_heading('2.8 Send Window', level=2)

add_para('Field name: Send Window')
add_para('Type: Two number inputs (start hour and end hour)')
add_para('Default: 8 (8:00 AM) to 21 (9:00 PM)')
add_para('Range: 0–23 (24-hour format)')

add_para(
    'The send window restricts when marketing messages (emails and SMS) can be delivered to subscribers. '
    'Messages scheduled outside this window will be held and sent at the start of the next window.'
)

add_para(
    'The hours are based on the subscriber\'s local time when available, or your site\'s timezone as a fallback. '
    'For example, with the default settings of 8 AM to 9 PM, no marketing messages will be sent between 9:01 PM and 7:59 AM.'
)

add_callout('IMPORTANT', 'The send window is especially important for SMS messages. '
            'Sending text messages in the middle of the night can lead to complaints and carrier filtering. '
            'Most SMS marketing regulations require messages to be sent during reasonable hours.')

add_callout('TIP', 'If your customer base is in a single timezone, the default 8 AM – 9 PM window works well. '
            'If you have a global audience, the plugin will attempt to use each subscriber\'s timezone when possible.')

# 2.9
doc.add_heading('2.9 Attribution Window', level=2)

add_para('Field name: Attribution Window (days)')
add_para('Type: Number input')
add_para('Default: 5 days')
add_para('Minimum: 1 day')

add_para(
    'The attribution window determines how long after a marketing message is sent (and clicked) the plugin will '
    'credit that message with revenue from a subsequent purchase. This uses a last-click attribution model.'
)

add_para(
    'For example, with the default setting of 5 days: if a subscriber clicks a link in your marketing email on Monday '
    'and makes a purchase on Wednesday (2 days later), that order\'s revenue is attributed to the email. '
    'If the same subscriber makes another purchase the following Monday (7 days later), that order is NOT attributed '
    'to the email because it falls outside the 5-day window.'
)

add_callout('TIP', 'A 5-day window is a good balance for most stores. Shorter windows (1–3 days) give you more conservative '
            'attribution numbers. Longer windows (7–14 days) capture more downstream purchases but may overcount.')

# 2.10
doc.add_heading('2.10 Unsubscribe Page Customisation', level=2)

doc.add_heading('2.10.1 Unsubscribe Page Title', level=3)

add_para('Field name: Unsubscribe Page Title')
add_para('Type: Text input')
add_para('Default: "Unsubscribe"')

add_para(
    'This is the HTML page title shown when a subscriber clicks an unsubscribe link. '
    'It appears in the browser tab and at the top of the unsubscribe confirmation page.'
)

doc.add_heading('2.10.2 Unsubscribe Message', level=3)

add_para('Field name: Unsubscribe Message')
add_para('Type: Textarea')
add_para('Default: "You have been successfully unsubscribed."')

add_para(
    'This is the message displayed to subscribers after they successfully unsubscribe. '
    'You can customise this to include a friendly message, a link back to your store, '
    'or information about re-subscribing.'
)

# 2.11
doc.add_heading('2.11 Uninstall Data Option', level=2)

add_para('Field name: Uninstall Data')
add_para('Type: Checkbox')
add_para('Default: Disabled (unchecked)')

add_para(
    'When enabled, all plugin data (all 14 database tables and all plugin settings) will be permanently deleted '
    'if you uninstall the plugin through the WordPress admin. This cannot be undone.'
)

add_callout('IMPORTANT', 'Only enable this option if you are completely sure you want to remove all data when uninstalling. '
            'If you are temporarily deactivating the plugin, leave this unchecked — deactivation does NOT delete data.')

# 2.12
doc.add_heading('2.12 Saving Your Settings', level=2)

add_steps([
    'After adjusting any settings on the page, scroll to the bottom of the form.',
    'Click the Save Settings button.',
    'A green success notice will appear at the top of the page confirming: "Settings saved."',
    'Your changes take effect immediately.',
])

add_callout('TIP', 'If you navigate away from the Settings page without clicking Save Settings, your changes will be lost. '
            'Always save before leaving the page.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: GDPR & COMPLIANCE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Chapter 3: GDPR & Compliance', level=1)

add_para(
    'Apotheca® Marketing Suite includes built-in tools to help you comply with the General Data Protection Regulation (GDPR) '
    'and other privacy regulations. This chapter explains how each compliance feature works.'
)

add_callout('IMPORTANT', 'While this plugin provides compliance tools, it is your responsibility to ensure that your overall '
            'marketing practices comply with applicable laws. Consult a legal professional if you are unsure about your obligations.')

# 3.1
doc.add_heading('3.1 Understanding Double Opt-In', level=2)

add_para(
    'Double opt-in is a two-step subscription process that provides verifiable proof that a subscriber genuinely '
    'requested to receive your marketing messages. It works as follows:'
)

add_bullet_list([
    'Single opt-in: The subscriber fills in their email and is immediately added as an active subscriber. This is simpler but offers less proof of consent.',
    'Double opt-in: The subscriber fills in their email, receives a confirmation email, and must click a link to confirm. Only after confirmation are they marked as an active subscriber.',
])

add_para(
    'Double opt-in is controlled by the Double Opt-In toggle on the Settings page (see Section 2.4). '
    'When enabled, it applies to all subscriber capture methods: checkout, registration, and forms.'
)

# 3.2
doc.add_heading('3.2 The Confirmation Email Flow', level=2)

add_para('When double opt-in is enabled, here is exactly what happens when a new subscriber is captured:')

add_steps([
    'A visitor opts in (via checkout, registration, or a form).',
    'The plugin creates a new subscriber record in the ams_subscribers table with status set to "pending".',
    'A unique unsubscribe token (48-character random string) is generated and saved to the subscriber record.',
    'The plugin fires the ams_send_double_optin_email action, which sends a confirmation email.',
    'The confirmation email is sent using wp_mail() and contains: a greeting using the subscriber\'s first name (or "there" if no name is provided), your site name, and a confirmation URL.',
    'The confirmation URL follows this format: https://yoursite.com/?ams_confirm=1&token=UNIQUE_TOKEN',
    'The subscriber clicks the link in the email.',
    'The plugin verifies the token, finds the matching subscriber record, and checks that the status is "pending".',
    'If valid, the subscriber\'s status is updated to "subscribed", the subscribed_at timestamp is recorded, '
    'gdpr_consent is set to 1 (true), and the gdpr_timestamp is recorded with the current date and time.',
    'The subscriber sees a confirmation page: "Your subscription has been confirmed. Thank you!"',
    'The ams_subscriber_confirmed action fires, which can trigger automation flows (e.g., a welcome series).',
])

add_callout('TIP', 'If a subscriber never clicks the confirmation link, their status remains "pending" and they '
            'will not receive any marketing emails. You may want to create a flow that sends a reminder after 24–48 hours.')

add_callout('IMPORTANT', 'The confirmation email is sent as a plain-text email via WordPress\'s wp_mail() function. '
            'Make sure your site\'s email sending is properly configured (via SMTP plugin or transactional email service) '
            'so the confirmation email does not end up in spam.')

# 3.3
doc.add_heading('3.3 The Unsubscribe Endpoint', level=2)

add_para(
    'The plugin registers a custom public URL endpoint for handling unsubscribe requests:'
)

p = doc.add_paragraph()
r = p.add_run('https://yoursite.com/ams-unsubscribe/?token=UNIQUE_TOKEN')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Consolas'

add_para(
    'This endpoint is created using WordPress rewrite rules. When the plugin is activated, it registers the '
    'rewrite rule ams-unsubscribe and flushes the permalink rules so the URL becomes active immediately.'
)

add_para('When a subscriber visits this URL with a valid token:')

add_steps([
    'The plugin intercepts the request on the template_redirect hook.',
    'It extracts the token parameter from the URL query string.',
    'It looks up the subscriber in the database using the token.',
    'If the subscriber is found and is not already unsubscribed, the plugin sets their status to "unsubscribed" '
    'and records the unsubscribed_at timestamp.',
    'The ams_subscriber_unsubscribed action fires (which you can use for custom integrations).',
    'The subscriber sees the unsubscribe confirmation page with the title and message you configured in Settings '
    '(see Section 2.10). By default, this shows "Unsubscribe" as the title and "You have been successfully unsubscribed." as the message.',
])

add_callout('IMPORTANT', 'If someone visits the unsubscribe URL with an invalid or missing token, they will see '
            'an error page: "Invalid unsubscribe link." This prevents unauthorized unsubscriptions.')

# 3.4
doc.add_heading('3.4 How Tokenised Unsubscribe Links Work', level=2)

add_para(
    'Every subscriber in Apotheca® Marketing Suite is assigned a unique unsubscribe token when their record is created. '
    'This token is a 48-character random string generated using WordPress\'s wp_generate_password() function '
    '(with special characters disabled for URL safety).'
)

add_para('The token is used in two ways:')

add_bullet_list([
    'Unsubscribe links: Every marketing email includes an unsubscribe link that contains the subscriber\'s token. '
    'When clicked, it takes the subscriber directly to the unsubscribe endpoint without requiring them to log in or enter their email.',
    'Confirmation links: When double opt-in is enabled, the confirmation email uses the same token to verify the subscriber\'s identity.',
])

add_para(
    'Because each token is unique to a subscriber and is 48 characters long, it is practically impossible for '
    'someone to guess another subscriber\'s token. This ensures that only the intended recipient can unsubscribe or confirm.'
)

add_callout('TIP', 'The unsubscribe token is stored in the unsubscribe_token column of the ams_subscribers table '
            'and is indexed for fast lookups.')

# 3.5
doc.add_heading('3.5 Consent Timestamp Tracking', level=2)

add_para(
    'For GDPR compliance, the plugin tracks exactly when each subscriber gave their consent. '
    'Two fields are maintained in each subscriber record:'
)

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
add_table_row_data(table4, [
    ['Field', 'Description'],
    ['gdpr_consent', 'A boolean flag (0 or 1) indicating whether the subscriber has given consent. '
     'Set to 1 when the subscriber opts in (or confirms via double opt-in).'],
    ['gdpr_timestamp', 'The exact date and time (in UTC) when consent was recorded. '
     'This provides an auditable timestamp that you can reference if a subscriber or regulator asks when consent was given.'],
])

add_para('')
add_para(
    'When a subscriber opts in at checkout with double opt-in disabled, the gdpr_consent flag and timestamp are '
    'set immediately at the time of checkout. When double opt-in is enabled, these fields are set only after '
    'the subscriber clicks the confirmation link — providing stronger proof of active consent.'
)

add_callout('IMPORTANT', 'Keep your database backed up regularly. The gdpr_consent and gdpr_timestamp fields are your '
            'proof of consent in the event of an audit.')

# 3.6
doc.add_heading('3.6 One-Click Unsubscribe for Subscribers', level=2)

add_para(
    'The unsubscribe process in Apotheca® Marketing Suite is designed to be as simple as possible for subscribers, '
    'in compliance with regulations like GDPR and CAN-SPAM:'
)

add_bullet_list([
    'Every marketing email includes a unique unsubscribe link at the bottom.',
    'Clicking the link immediately unsubscribes the subscriber — no login required, no confirmation form, no additional clicks.',
    'The subscriber sees a friendly confirmation page with your customised message.',
    'The subscriber\'s status is updated to "unsubscribed" in the database.',
    'The unsubscribed_at timestamp is recorded for your records.',
    'The subscriber will not receive any further marketing emails or SMS messages from the plugin.',
])

add_callout('TIP', 'If a subscriber unsubscribes by mistake, they cannot re-subscribe through the same link. '
            'They would need to opt in again through your checkout, registration, or a sign-up form. '
            'The plugin will then create a new consent record with a fresh timestamp.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: SUBSCRIBER MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Chapter 4: Subscriber Management', level=1)

add_para(
    'This chapter covers how subscribers are captured, stored, and managed within Apotheca® Marketing Suite. '
    'The Subscribers page is one of the most frequently used areas of the plugin.'
)

# 4.1
doc.add_heading('4.1 How Subscribers Are Captured', level=2)

add_para(
    'Subscribers can enter your list through several channels. Each method records a different source value '
    'so you can track where your subscribers come from.'
)

doc.add_heading('4.1.1 Checkout Capture', level=3)

add_para(
    'When the Checkout Opt-In setting is enabled (see Section 2.2), a marketing consent checkbox appears on '
    'your WooCommerce checkout page just above the "Place Order" button. Here is what happens:'
)

add_steps([
    'The customer fills in their billing details and reaches the checkout form.',
    'They see a checkbox with your configured label text (e.g., "Keep me updated with news and offers via email.").',
    'If the customer checks the box and completes the order, the plugin captures their email address, first name, '
    'last name, and phone number from the billing fields.',
    'A subscriber record is created (or updated if one already exists for that email address) with source set to "checkout".',
    'If double opt-in is disabled, the subscriber is immediately marked as "subscribed."',
    'If double opt-in is enabled, the subscriber is marked as "pending" and a confirmation email is sent.',
    'If the customer does NOT check the box, a subscriber record is still created but with status "never_subscribed." '
    'This allows you to see checkout customers who declined marketing, without sending them any messages.',
])

doc.add_heading('4.1.2 Registration Capture', level=3)

add_para(
    'When the Registration Capture setting is enabled (see Section 2.3), new WooCommerce account registrations '
    'are automatically added to your subscriber list.'
)

add_steps([
    'A visitor creates a new account on your site (via the My Account page or during checkout).',
    'The plugin captures their email address, first name, and last name from their WordPress user profile.',
    'A subscriber record is created with source set to "registration".',
    'If double opt-in is enabled, a confirmation email is sent and the subscriber starts as "pending."',
    'If double opt-in is disabled, the subscriber is immediately marked as "subscribed."',
])

doc.add_heading('4.1.3 Form Capture', level=3)

add_para(
    'Pop-up forms (covered in a later chapter) can also capture subscribers. When a visitor fills in a pop-up '
    'form on your site, their data is submitted to the plugin and a subscriber record is created. '
    'The source for form-captured subscribers reflects the form submission method.'
)

# 4.2
doc.add_heading('4.2 Viewing the Subscriber List', level=2)

add_steps([
    'In the WordPress admin sidebar, click Apotheca® Marketing → Subscribers.',
    'The Subscribers page displays a table listing all subscribers, 20 per page.',
    'The table shows the following columns: Email, Name, Status, Source, Orders, Total Spent, RFM, and Subscribed (date).',
    'If you have more than 20 subscribers, pagination links appear at the bottom of the table. '
    'Click the page numbers to navigate through your subscriber list.',
])

# 4.3
doc.add_heading('4.3 Searching and Filtering Subscribers', level=2)

doc.add_heading('4.3.1 Searching', level=3)

add_steps([
    'On the Subscribers page, locate the search box at the top of the page.',
    'Type a search term in the text field. You can search by email address, first name, or last name.',
    'Click the Search button (or press Enter).',
    'The table updates to show only subscribers matching your search. The search performs a partial match '
    '(e.g., searching for "john" will find "john@example.com", "Johnny Smith", and "Sarah Johnson").',
    'To clear the search, delete the text in the search box and click Search again.',
])

doc.add_heading('4.3.2 Filtering by Status', level=3)

add_para(
    'You can filter the subscriber list by status using the status query parameter in the URL. '
    'For example, adding ?status=subscribed to the Subscribers page URL will show only active subscribers.'
)

add_callout('TIP', 'You can combine search and status filtering. For example, search for "gmail" with a status filter of '
            '"subscribed" to see all active Gmail subscribers.')

# 4.4
doc.add_heading('4.4 Understanding the Subscriber Table Columns', level=2)

table5 = doc.add_table(rows=1, cols=2)
table5.style = 'Table Grid'
add_table_row_data(table5, [
    ['Column', 'Description'],
    ['Email', 'The subscriber\'s email address. Displayed in bold for easy scanning.'],
    ['Name', 'The subscriber\'s full name (first name + last name). May be blank if no name was captured.'],
    ['Status', 'The subscriber\'s current status, displayed as a coloured badge (see Section 4.5).'],
    ['Source', 'How the subscriber was captured: "Checkout", "Registration", or other sources. Displayed with the first letter capitalised.'],
    ['Orders', 'The total number of WooCommerce orders placed by this subscriber.'],
    ['Total Spent', 'The total revenue from all orders placed by this subscriber, formatted in your store\'s currency.'],
    ['RFM', 'The subscriber\'s RFM (Recency, Frequency, Monetary) segment label, such as "Champion", "At Risk", etc. Shows a dash (—) if not yet calculated.'],
    ['Subscribed', 'The date the subscriber became active, formatted according to your WordPress date format setting. Shows a dash (—) if the subscriber has not yet confirmed.'],
])

# 4.5
doc.add_heading('4.5 Subscriber Statuses Explained', level=2)

add_para('Each subscriber has one of the following statuses:')

table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Table Grid'
add_table_row_data(table6, [
    ['Status', 'Meaning'],
    ['subscribed', 'The subscriber is active and will receive marketing messages. This is the status after opting in (single opt-in) or confirming (double opt-in).'],
    ['pending', 'The subscriber has opted in but has not yet confirmed via the double opt-in email. They will NOT receive marketing messages until they confirm.'],
    ['unsubscribed', 'The subscriber has opted out of marketing messages by clicking an unsubscribe link. They will not receive any further messages.'],
    ['never_subscribed', 'The subscriber placed an order at checkout but did NOT check the marketing consent box. Their data is stored for order tracking, but they will not receive marketing messages.'],
])

add_para('')
add_callout('TIP', 'You can use segments (covered in a later chapter) to target subscribers by status. '
            'For example, you could create a segment of "pending" subscribers to track how many have not yet confirmed.')

# 4.6
doc.add_heading('4.6 Subscriber Sources', level=2)

add_para('The source field tells you how each subscriber entered your list:')

table7 = doc.add_table(rows=1, cols=2)
table7.style = 'Table Grid'
add_table_row_data(table7, [
    ['Source', 'How the Subscriber Was Captured'],
    ['checkout', 'Opted in via the marketing consent checkbox on the WooCommerce checkout page.'],
    ['registration', 'Automatically captured when they created a new WooCommerce/WordPress account.'],
    ['form', 'Submitted a pop-up form created with the Forms feature (covered in a later chapter).'],
    ['import', 'Added via CSV import (see Section 4.9).'],
    ['manual', 'Manually added by an administrator.'],
])

# 4.7
doc.add_heading('4.7 Tags', level=2)

add_para(
    'Tags are flexible labels you can attach to any subscriber. They are stored as a JSON array in the '
    'tags column of the subscriber record. Tags can be used for:'
)

add_bullet_list([
    'Organising subscribers into custom groups (e.g., "VIP", "wholesale", "holiday-promo-2024").',
    'Targeting specific groups in segments and campaigns.',
    'Triggering or filtering automation flows based on tag presence.',
])

add_callout('TIP', 'Tags are freeform text — you can create any tag name you like. Use consistent naming conventions '
            '(e.g., lowercase with hyphens) to keep your tags organised.')

# 4.8
doc.add_heading('4.8 RFM Score and Segment', level=2)

add_para(
    'Every subscriber with at least one order is automatically scored using the RFM (Recency, Frequency, Monetary) model. '
    'This is calculated nightly by the plugin\'s built-in analytics engine.'
)

doc.add_heading('4.8.1 RFM Score', level=3)

add_para(
    'The RFM score is a short string (up to 5 characters) that represents the subscriber\'s purchasing behaviour. '
    'It is stored in the rfm_score column.'
)

doc.add_heading('4.8.2 RFM Segment', level=3)

add_para(
    'Based on the RFM score, each subscriber is assigned to an RFM segment (stored in rfm_segment). '
    'Common segment labels include:'
)

add_bullet_list([
    'Champion — Recent, frequent, high-spending customers.',
    'Loyal — Frequent buyers who spend well.',
    'Potential Loyalist — Recent customers with growing order frequency.',
    'At Risk — Previously active customers who have not ordered recently.',
    'Hibernating — Customers with very old, infrequent purchases.',
    'Lost — Customers who have not purchased in a very long time.',
])

add_callout('TIP', 'Use RFM segments to power your automation flows. For example, create a win-back flow that targets '
            '"At Risk" and "Hibernating" subscribers with special offers.')

# 4.9
doc.add_heading('4.9 Subscriber Profile Page', level=2)

add_para(
    'Each subscriber in the Apotheca® Marketing Suite database has a rich profile that stores detailed information '
    'about their interactions with your store. The subscriber profile includes the following data points:'
)

doc.add_heading('4.9.1 Core Information', level=3)

add_bullet_list([
    'Email address',
    'Phone number (if provided)',
    'First and last name',
    'Subscription status and source',
    'Subscription date and unsubscription date (if applicable)',
])

doc.add_heading('4.9.2 GDPR Consent Information', level=3)

add_bullet_list([
    'Whether consent was given (gdpr_consent flag)',
    'Exact timestamp of when consent was recorded (gdpr_timestamp)',
])

doc.add_heading('4.9.3 Purchase History', level=3)

add_bullet_list([
    'Total number of orders',
    'Total amount spent',
    'Date of last order',
])

doc.add_heading('4.9.4 RFM & Predictive Data', level=3)

add_bullet_list([
    'RFM score and segment',
    'Predicted customer lifetime value (CLV)',
    'Predicted next order date',
    'Churn risk score (0–255, where higher means more likely to churn)',
])

doc.add_heading('4.9.5 Communication Preferences', level=3)

add_bullet_list([
    'SMS opt-in status',
    'Best send hour (optimised by the AI engine, default is 10 AM)',
])

doc.add_heading('4.9.6 Event Timeline', level=3)

add_para(
    'The ams_events table stores a complete timeline of each subscriber\'s activity. Events include orders placed, '
    'emails opened, links clicked, and other interactions. Each event records the event type, associated data, '
    'WooCommerce order ID (if applicable), related product IDs, and a timestamp.'
)

doc.add_heading('4.9.7 Send History', level=3)

add_para(
    'The ams_sends table tracks every email and SMS sent to the subscriber, including: which campaign or flow step '
    'triggered the send, the channel (email or SMS), delivery status, timestamps for when the message was sent, opened, '
    'clicked, bounced, or resulted in an unsubscribe, and any revenue attributed to the message.'
)

doc.add_heading('4.9.8 Custom Fields and Tags', level=3)

add_para(
    'Both tags and custom_fields are stored as JSON in the subscriber record, giving you flexible extensibility '
    'for any additional data you want to track per subscriber.'
)

add_callout('TIP', 'The subscriber profile is a powerful tool for understanding individual customer behaviour. '
            'Use it to investigate why a particular customer may have stopped purchasing, or to verify that an automation '
            'flow sent the correct messages.')

# ── Footer ────────────────────────────────────────────────────────────────────

doc.add_page_break()

doc.add_heading('End of Part 1', level=1)

add_para(
    'This concludes Part 1 of the Apotheca® Marketing Suite User Manual, covering Installation & Activation, '
    'Global Settings & Configuration, GDPR & Compliance, and Subscriber Management.'
)

add_para(
    'Part 2 will cover: Automation Flows, Campaigns, Segments, Pop-Up Forms, SMS, Analytics, '
    'AI Features, Reviews, and Multi-Site Sync.'
)

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Apotheca_User_Manual_Part1.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
