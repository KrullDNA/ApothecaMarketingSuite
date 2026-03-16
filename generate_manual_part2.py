#!/usr/bin/env python3
"""Generate Apotheca Marketing Suite User Manual — Part 2 (Chapters 5–10)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# Helper functions
def add_heading1(text):
    p = doc.add_heading(text, level=1)
    doc.add_paragraph()  # spacer
    return p

def add_heading2(text):
    return doc.add_heading(text, level=2)

def add_heading3(text):
    return doc.add_heading(text, level=3)

def add_para(text):
    return doc.add_paragraph(text)

def add_bold_para(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(text)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x2F)
    p.add_run(text)
    return p

def add_important(text):
    p = doc.add_paragraph()
    run = p.add_run('IMPORTANT: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(text)
    return p

def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {number}: ')
    run.bold = True
    p.add_run(text)
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_bold_bullet(label, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    p.add_run(text)
    return p


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: AUTOMATED FLOWS — OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 5: Automated Flows — Overview')

add_heading2('5.1 What Are Flows?')
add_para(
    'Flows are automated sequences of actions that run when a specific event (called a trigger) occurs. '
    'Think of a flow as a "set it and forget it" automation: once you build and activate a flow, '
    'the Apotheca Marketing Suite handles everything in the background — sending emails, sending SMS messages, '
    'adding tags, waiting specified periods, branching based on conditions, and more.'
)
add_para(
    'Each flow consists of three parts:'
)
add_bold_bullet('Trigger — ', 'the event that starts the flow (e.g., a new subscriber signs up, a cart is abandoned).')
add_bold_bullet('Steps — ', 'the ordered sequence of actions the flow performs (e.g., wait 1 hour, send an email, add a tag).')
add_bold_bullet('Enrolments — ', 'individual subscriber journeys through the flow. When a trigger fires for a subscriber, they are "enrolled" into the flow and progress through each step automatically.')

add_para(
    'The plugin supports 8 trigger types and 8 step types, giving you enormous flexibility to build '
    'automations for welcome series, abandoned cart recovery, post-purchase follow-ups, win-back campaigns, and more.'
)

add_heading2('5.2 Navigating to the Flows Page')
add_step(1, 'In your WordPress admin dashboard, look for the Apotheca Marketing menu in the left sidebar.')
add_step(2, 'Click Flows to open the Flows management page.')
add_step(3, 'You will see a list of all existing flows. If this is your first time, the list will be empty.')

add_heading2('5.3 The Flow Builder Interface')
add_para(
    'The flow builder is a visual, drag-and-drop interface built on React. It lets you design your '
    'automation sequence without writing any code.'
)
add_para('The flow builder has three main areas:')
add_bold_bullet('Flow List — ', 'shows all your existing flows in a table. You can see each flow\'s name, trigger type, status (active/inactive), and the date it was created. This is also where you can import pre-built templates.')
add_bold_bullet('Flow Editor — ', 'opens when you create or edit a flow. Here you configure the trigger, add steps in sequence, and drag to reorder them. An unsaved changes indicator appears whenever you have modifications that haven\'t been saved yet.')
add_bold_bullet('Step Editor — ', 'a context-sensitive panel that appears when you click on a step. It shows the configuration options specific to that step type (e.g., email subject and body for a Send Email step, delay duration for a Wait step).')

add_heading2('5.4 Creating a New Flow')
add_step(1, 'On the Flows page, click the Create New Flow button.')
add_step(2, 'Enter a descriptive name for your flow in the Flow Name field (e.g., "Welcome Series" or "Abandoned Cart Recovery").')
add_step(3, 'Select a Trigger Type from the dropdown. This determines what event will start the flow. See Chapter 6 for details on each trigger type.')
add_step(4, 'If your chosen trigger has additional configuration options (e.g., the Win-Back trigger requires you to set the number of days since last order), fill those in now.')
add_step(5, 'Click Add Step to begin building your automation sequence. See Chapter 7 for details on each step type.')
add_step(6, 'Continue adding steps to build your complete automation.')
add_step(7, 'Click Save Flow to save your work.')

add_tip('You can drag and drop steps to reorder them at any time. The flow builder will show an unsaved changes indicator when you have modifications that need saving.')

add_heading2('5.5 Naming and Saving Flows')
add_para(
    'Use clear, descriptive names for your flows so you can identify them at a glance in the flow list. '
    'Good examples: "Welcome Series — 3 Emails", "Cart Recovery with Discount", "VIP Win-Back Campaign".'
)
add_para(
    'The flow builder communicates with the WordPress REST API (at the /ams/v1 endpoint) to save your flows. '
    'When you click Save Flow, all steps and their configurations are saved at once. '
    'The builder replaces all existing steps with your current configuration each time you save — there is no partial saving.'
)
add_important('Always click Save Flow before leaving the page. If you see the unsaved changes indicator, your changes have not been saved yet.')

add_heading2('5.6 Activating and Deactivating Flows')
add_step(1, 'In the flow list, locate the flow you want to activate or deactivate.')
add_step(2, 'Toggle the Status switch to Active to start the flow, or Inactive to pause it.')
add_step(3, 'When a flow is Active, its trigger is live — any matching event will enrol subscribers into the flow.')
add_step(4, 'When a flow is Inactive, the trigger stops firing. Subscribers who are already enrolled will continue progressing through their remaining steps, but no new subscribers will be enrolled.')

add_important('Deactivating a flow does not cancel enrolments that are already in progress. It only prevents new enrolments.')

add_heading2('5.7 How Flow Enrolment Works')
add_para(
    'When a trigger fires for a subscriber, the system creates an enrolment record that tracks the subscriber\'s '
    'progress through the flow. Here is how enrolment works:'
)
add_bold_bullet('Deduplication — ', 'a subscriber cannot be enrolled in the same flow more than once at the same time. If a subscriber is already actively progressing through a flow and the trigger fires again, the second enrolment is skipped. This prevents duplicate messages.')
add_bold_bullet('Enrolment status — ', 'each enrolment has a status: active (currently progressing through steps), completed (finished all steps), or exited (removed early, e.g., due to unsubscribing or hitting an Exit Flow step).')
add_bold_bullet('Automatic exit on unsubscribe — ', 'if a subscriber unsubscribes, ALL of their active flow enrolments are automatically exited. The system listens for the unsubscribe event and exits every active enrolment for that subscriber.')
add_bold_bullet('Step progression — ', 'after each step completes, the system automatically advances the subscriber to the next step. For Wait steps, the next step is scheduled to run after the configured delay using Action Scheduler.')

add_heading2('5.8 Frequency Caps and Send Windows')
add_para(
    'To protect your subscribers from being overwhelmed with messages, Apotheca Marketing Suite enforces '
    'frequency caps and send windows on all flow-triggered communications.'
)

add_heading3('Frequency Caps')
add_para('Frequency caps limit the number of messages a subscriber can receive in a 24-hour period:')
add_bold_bullet('Email: ', '3 emails per subscriber per 24-hour period (default). You can adjust this in the plugin settings under the frequency_cap_email option.')
add_bold_bullet('SMS: ', '2 SMS messages per subscriber per 24-hour period (default). You can adjust this in the plugin settings under the frequency_cap_sms option.')
add_para(
    'When a flow tries to send a message but the subscriber has already hit the frequency cap, '
    'the step is automatically rescheduled to retry 1 hour later. This continues until the cap resets.'
)

add_heading3('Send Windows')
add_para(
    'Send windows ensure that messages are only delivered during reasonable hours in the subscriber\'s local timezone:'
)
add_bold_bullet('Default window: ', '8:00 AM to 9:00 PM in the subscriber\'s local time.')
add_bold_bullet('Timezone detection: ', 'the plugin determines the subscriber\'s timezone from their WooCommerce billing country. It supports over 35 countries with automatic timezone mapping (e.g., US → America/New_York, GB → Europe/London, AU → Australia/Sydney). If no billing country is available, the site\'s WordPress timezone is used as a fallback.')
add_bold_bullet('Rescheduling: ', 'if a message would be sent outside the send window, it is automatically rescheduled to the next window opening (e.g., 8:00 AM the following day in the subscriber\'s timezone).')
add_para('You can adjust the send window start and end hours in the plugin settings under send_window_start and send_window_end.')

add_tip('Frequency caps and send windows only apply to Send Email and Send SMS steps. Other step types (Add Tag, Wait, Condition Branch, etc.) execute immediately regardless of the time of day.')


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: FLOW TRIGGERS (8 TYPES)
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 6: Flow Triggers — Step by Step for Each')

add_para(
    'A trigger is the event that starts a flow. When the trigger condition is met for a subscriber, '
    'they are automatically enrolled into the flow and begin progressing through the steps. '
    'Apotheca Marketing Suite provides 8 trigger types, each designed for a specific marketing scenario.'
)

# ── 6.1 Welcome Series ──
add_heading2('6.1 Welcome Series Trigger')

add_heading3('What It Does')
add_para(
    'The Welcome Series trigger fires when a new subscriber confirms their subscription. It is the '
    'ideal trigger for onboarding new subscribers with a series of introductory emails.'
)

add_heading3('When It Fires')
add_para('The trigger fires in one of three scenarios:')
add_bold_bullet('Double opt-in enabled: ', 'fires when the subscriber confirms their email address (i.e., clicks the confirmation link). The plugin listens for the ams_subscriber_confirmed event.')
add_bold_bullet('No double opt-in — checkout: ', 'fires when a customer completes a WooCommerce checkout and their subscription status is "subscribed". This fires on the woocommerce_checkout_order_processed event.')
add_bold_bullet('No double opt-in — registration: ', 'fires when a new WordPress user registers and their subscription status is "subscribed". This fires on the user_register event.')

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Welcome Series".')
add_step(3, 'From the Trigger Type dropdown, select Welcome Series.')
add_step(4, 'No additional trigger configuration is required for this trigger type.')
add_step(5, 'Add your welcome steps (e.g., Send Email immediately, Wait 2 days, Send Email with brand story, etc.).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_para('This trigger has no additional configuration options. It automatically fires on subscriber confirmation or direct subscription.')

add_heading3('Example Use Case')
add_para(
    'Create a 5-email welcome series: send a welcome email immediately, wait 2 days and share your brand story, '
    'wait 3 days and showcase trending products, wait 2 days and offer a 10% first-purchase discount, '
    'then wait 3 days and check if a purchase was made with a condition branch.'
)

# ── 6.2 Abandoned Cart ──
add_heading2('6.2 Abandoned Cart Trigger')

add_heading3('What It Does')
add_para(
    'The Abandoned Cart trigger fires when a subscriber\'s shopping cart is detected as abandoned. '
    'This is one of the highest-ROI automations you can create, recovering revenue from shoppers who left without completing their purchase.'
)

add_heading3('When It Fires')
add_para(
    'The trigger fires when the Apotheca abandoned cart detector flags a cart as abandoned. '
    'The system tracks cart activity and fires the ams_cart_abandoned event when a subscriber adds items to their cart '
    'but does not complete checkout within the configured abandonment window.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Abandoned Cart Recovery".')
add_step(3, 'From the Trigger Type dropdown, select Abandoned Cart.')
add_step(4, 'No additional trigger configuration is required — the abandonment detection timing is configured in the plugin\'s cart tracking settings.')
add_step(5, 'Add your recovery steps. A typical sequence: Wait 1 hour → Send Email (reminder) → Wait 23 hours → Send Email (social proof) → Wait 24 hours → Send Email (discount offer).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_para('This trigger has no additional configuration options on the trigger itself. Cart abandonment timing is controlled by the cart tracking module.')

add_heading3('Example Use Case')
add_para(
    'Build a 3-email recovery sequence: (1) a gentle reminder 1 hour after abandonment, '
    '(2) a social proof email with customer reviews 24 hours later, '
    '(3) a final email with a 10% discount code (e.g., CART10) 48 hours after abandonment.'
)

# ── 6.3 Post-Purchase ──
add_heading2('6.3 Post-Purchase Trigger')

add_heading3('What It Does')
add_para(
    'The Post-Purchase trigger fires when a WooCommerce order is marked as completed. '
    'Use it to send thank-you emails, request product reviews, cross-sell related products, or check in on customer satisfaction.'
)

add_heading3('When It Fires')
add_para(
    'The trigger fires on the woocommerce_order_status_completed event. It looks up the billing email '
    'from the order, finds the matching subscriber, and enrols them into the flow — but only if the subscriber\'s '
    'status is "subscribed".'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Post-Purchase Follow-Up".')
add_step(3, 'From the Trigger Type dropdown, select Post-Purchase.')
add_step(4, 'No additional trigger configuration is required.')
add_step(5, 'Add your post-purchase steps. A common sequence: Send Email (thank you) immediately → Wait 7 days → Send Email (review request).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_para('This trigger has no additional configuration options. It fires automatically on every completed order for subscribed customers.')

add_heading3('Example Use Case')
add_para(
    'Send an immediate thank-you email after purchase, wait 7 days, then send a review request email '
    'with a star rating widget (1–5 stars). This helps build social proof and customer engagement.'
)

# ── 6.4 Win-Back ──
add_heading2('6.4 Win-Back Trigger')

add_heading3('What It Does')
add_para(
    'The Win-Back trigger identifies subscribers who have not placed an order in a specified number of days '
    'and automatically enrols them into a re-engagement flow. This is essential for reducing customer churn.'
)

add_heading3('When It Fires')
add_para(
    'Unlike event-based triggers, the Win-Back trigger runs on a daily schedule via Action Scheduler. '
    'Every day, it queries the subscriber database for customers whose last order date is older than the configured threshold. '
    'It processes up to 200 qualifying subscribers per run.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Win-Back Campaign — 90 Days".')
add_step(3, 'From the Trigger Type dropdown, select Win-Back.')
add_step(4, 'In the Days Since Last Order field, enter the number of days of inactivity required before the trigger fires. The default is 90 days.')
add_step(5, 'Add your win-back steps. A typical sequence: Send Email (we miss you) → Wait 15 days → Send Email (discount offer) → Wait 15 days → Send SMS (final reminder).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_bold_bullet('Days Since Last Order — ', 'the number of days since the subscriber\'s last completed order. Only subscribers with at least 1 previous order and a last order date older than this threshold will be enrolled. Default: 90 days.')

add_heading3('Example Use Case')
add_para(
    'Target customers who haven\'t purchased in 90+ days with a 3-touch sequence: '
    'a "We miss you" email on day 0, a 15% discount offer (code: COMEBACK15) on day 15, '
    'and a final SMS reminder on day 30.'
)

add_tip('You can create multiple Win-Back flows with different thresholds — e.g., one at 60 days (gentle nudge) and another at 120 days (aggressive discount).')

# ── 6.5 Browse Abandonment ──
add_heading2('6.5 Browse Abandonment Trigger')

add_heading3('What It Does')
add_para(
    'The Browse Abandonment trigger fires when a subscriber views a product but does not add it to their cart '
    'within 30 minutes. It allows you to follow up with subscribers who showed interest but did not take action.'
)

add_heading3('When It Fires')
add_para(
    'This trigger runs on an hourly schedule via Action Scheduler. Each hour, it checks the event log '
    'for viewed_product events that are older than 30 minutes and have no subsequent added_to_cart event. '
    'To prevent duplicate triggers, the system records a browse_abandonment_triggered event after enrolling a subscriber, '
    'ensuring the same browse session is not triggered again. It processes up to 100 qualifying subscribers per run.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Browse Abandonment Follow-Up".')
add_step(3, 'From the Trigger Type dropdown, select Browse Abandonment.')
add_step(4, 'No additional trigger configuration is required.')
add_step(5, 'Add your follow-up steps. A typical sequence: Wait 30 minutes → Send Email (product reminder) → Wait 24 hours → Send Email (curated recommendations).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_para('This trigger has no additional configuration options. The 30-minute inactivity threshold and hourly check frequency are built in.')

add_heading3('Example Use Case')
add_para(
    'When a subscriber browses a product but doesn\'t add it to their cart, wait 30 minutes and send a reminder '
    'email featuring the product they viewed. If they still don\'t purchase, send a follow-up 24 hours later '
    'with curated product recommendations.'
)

add_important('Browse abandonment tracking requires that your WooCommerce product pages have the Apotheca tracking script installed. This is included automatically when the plugin is active.')

# ── 6.6 Birthday ──
add_heading2('6.6 Birthday Trigger')

add_heading3('What It Does')
add_para(
    'The Birthday trigger fires on a subscriber\'s birthday, allowing you to send personalised birthday greetings, '
    'special discounts, or exclusive offers.'
)

add_heading3('When It Fires')
add_para(
    'This trigger runs daily via Action Scheduler. Each day, it queries the subscriber database for subscribers '
    'whose birthday matches today\'s date (month and day). The birthday is stored in the subscriber\'s custom fields '
    'as either a single "birthday" field in MM-DD format, or as separate "birthday_month" (MM) and "birthday_day" (DD) fields. '
    'It processes up to 200 matching subscribers per run.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Birthday Celebration".')
add_step(3, 'From the Trigger Type dropdown, select Birthday.')
add_step(4, 'No additional trigger configuration is required.')
add_step(5, 'Add your birthday steps. For example: Send Email (birthday greeting with discount code) → Wait 7 days → Send Email (reminder to use birthday discount).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_para('This trigger has no additional configuration options. It matches any subscriber whose birthday custom field matches today\'s date.')

add_heading3('Example Use Case')
add_para(
    'Send a personalised "Happy Birthday!" email with a 20% discount code on the subscriber\'s birthday, '
    'then follow up 7 days later with a reminder to use the code before it expires.'
)

add_important(
    'For this trigger to work, subscribers must have a birthday stored in their custom fields. '
    'You can collect birthday data via your sign-up forms, import it from a CSV, or use the Update Field step in another flow. '
    'The birthday must be in MM-DD format (e.g., "03-15" for March 15th), or stored as separate birthday_month and birthday_day fields.'
)

# ── 6.7 RFM Change ──
add_heading2('6.7 RFM Change Trigger')

add_heading3('What It Does')
add_para(
    'The RFM Change trigger fires when a subscriber\'s RFM (Recency, Frequency, Monetary) segment changes. '
    'This allows you to react automatically when a customer\'s behaviour shifts — for example, when a "Champion" '
    'customer becomes "At Risk", or when a "Potential" customer upgrades to "Loyal".'
)

add_heading3('When It Fires')
add_para(
    'The trigger fires whenever the RFM scoring engine (which runs nightly) detects that a subscriber\'s RFM segment '
    'has changed from one named segment to another. The engine fires the ams_rfm_segment_changed event with the '
    'subscriber ID, old segment name, and new segment name.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "At Risk Customer Rescue".')
add_step(3, 'From the Trigger Type dropdown, select RFM Change.')
add_step(4, 'Optionally configure the From Segment and/or To Segment fields to limit which segment transitions trigger the flow. If left blank, any segment change will trigger the flow.')
add_step(5, 'Add your response steps (e.g., Send Email with a special offer to at-risk customers).')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_bold_bullet('From Segment — ', 'the RFM segment the subscriber is moving from. Leave blank to match any origin segment. Available segments: Champions, Big Spenders, Loyal, New Customers, Potential, At Risk, About to Sleep, Lost.')
add_bold_bullet('To Segment — ', 'the RFM segment the subscriber is moving to. Leave blank to match any destination segment. Available segments are the same as above.')

add_heading3('Example Use Case')
add_para(
    'Create a "Champion Rescue" flow that triggers when a subscriber moves from "Champions" to "At Risk". '
    'Send an exclusive VIP email with a personalised offer to re-engage them before they become a lost customer.'
)

add_tip('You can create separate flows for different segment transitions. For example, one flow for "At Risk → Lost" (aggressive win-back) and another for "Potential → Loyal" (congratulations and upsell).')

# ── 6.8 Custom Event ──
add_heading2('6.8 Custom Event Trigger')

add_heading3('What It Does')
add_para(
    'The Custom Event trigger fires when a custom event is recorded for a subscriber via the Apotheca event tracking system. '
    'This is the most flexible trigger type — it lets you create automations for any event you define, '
    'such as "completed_onboarding", "left_review", "referred_a_friend", or any other custom event you track.'
)

add_heading3('When It Fires')
add_para(
    'The trigger listens for the ams_event_recorded event. It fires for any event type that is NOT one of the '
    'standard built-in WooCommerce events (placed_order, completed_purchase, refund_requested, viewed_product, '
    'added_to_cart, started_checkout, abandoned_cart, wrote_review, browse_abandonment_triggered). '
    'Any event type not in that list is considered a "custom event" and will activate this trigger.'
)

add_heading3('How to Set It Up')
add_step(1, 'Navigate to Flows and click Create New Flow.')
add_step(2, 'Enter a name such as "Referral Thank-You".')
add_step(3, 'From the Trigger Type dropdown, select Custom Event.')
add_step(4, 'In the Event Type field, enter the exact event type string that will trigger this flow (e.g., "referred_a_friend"). If left blank, any custom event will trigger the flow.')
add_step(5, 'Add your response steps.')
add_step(6, 'Click Save Flow and set the status to Active.')

add_heading3('Configuration Options')
add_bold_bullet('Event Type — ', 'the specific custom event type to listen for. Must match exactly the event type string recorded by the EventTracker. Leave blank to trigger on any custom event (not recommended for production use).')

add_heading3('Example Use Case')
add_para(
    'Track a "referred_a_friend" custom event via the Apotheca event API. Create a flow that triggers on this event '
    'and sends a thank-you email with a referral bonus discount code. You could also add a tag "referrer" '
    'to the subscriber for future segmentation.'
)

add_important('Custom events must be recorded via the Apotheca event tracking API (EventTracker class) for this trigger to work. The event type string must match exactly — it is case-sensitive.')


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 7: FLOW STEPS (8 TYPES)
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 7: Flow Steps — Step by Step for Each')

add_para(
    'Steps are the building blocks of your flows. Each step performs a specific action — sending an email, '
    'waiting a set period, adding a tag, and so on. You chain steps together to create your automation sequence. '
    'Apotheca Marketing Suite provides 8 step types.'
)

# ── 7.1 Send Email ──
add_heading2('7.1 Send Email Step')

add_heading3('What It Does')
add_para(
    'The Send Email step sends a personalised email to the subscriber. It supports dynamic personalisation tokens, '
    'HTML content, preview text, and automatically appends an unsubscribe link for GDPR compliance.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Send Email from the step type dropdown.')
add_step(3, 'Fill in the Subject line.')
add_step(4, 'Write your email Body using the HTML editor.')
add_step(5, 'Optionally add Preview Text (the text shown in email client previews).')
add_step(6, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Subject — ', 'the email subject line. Supports personalisation tokens.')
add_bold_bullet('Body — ', 'the HTML email body. Supports personalisation tokens.')
add_bold_bullet('Preview Text — ', 'optional text that appears in email client preview panes. Rendered as a hidden div at the top of the email.')

add_para('Available personalisation tokens:')
add_bullet('{{first_name}} — subscriber\'s first name')
add_bullet('{{last_name}} — subscriber\'s last name')
add_bullet('{{email}} — subscriber\'s email address')
add_bullet('{{full_name}} — subscriber\'s full name (first + last)')
add_bullet('{{site_name}} — your WordPress site name')
add_bullet('{{site_url}} — your WordPress site URL')
add_bullet('{{unsubscribe_url}} — one-click unsubscribe link')
add_bullet('{{total_orders}} — subscriber\'s total number of orders')
add_bullet('{{total_spent}} — subscriber\'s total amount spent')
add_bullet('{{rfm_segment}} — subscriber\'s current RFM segment name')

add_heading3('Example Use Case')
add_para(
    'In an abandoned cart flow, send an email with the subject "Still thinking about it, {{first_name}}?" '
    'and a body that includes product images, a call-to-action button linking back to the cart, and a discount code.'
)

add_important('Every email automatically includes an unsubscribe link in the footer for GDPR/CAN-SPAM compliance. Do not remove or hide this link.')

# ── 7.2 Send SMS ──
add_heading2('7.2 Send SMS Step')

add_heading3('What It Does')
add_para(
    'The Send SMS step sends a text message (SMS) or multimedia message (MMS) to the subscriber. '
    'It validates that the subscriber has a phone number and has opted in to SMS before sending. '
    'A STOP instruction is automatically appended for TCPA compliance.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Send SMS from the step type dropdown.')
add_step(3, 'Write your SMS Message. Keep it concise — SMS messages have character limits.')
add_step(4, 'Optionally add a Media URL for MMS (e.g., a product image URL).')
add_step(5, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Message — ', 'the SMS text content. Supports the same personalisation tokens as email.')
add_bold_bullet('Media URL — ', 'optional URL of an image or media file to include as MMS.')

add_heading3('Example Use Case')
add_para(
    'At the end of a win-back flow, send an SMS saying: "Hi {{first_name}}, we have a special 15% off offer just for you! '
    'Use code COMEBACK15 at checkout. Shop now: {{site_url}}"'
)

add_important('SMS messages are queued asynchronously via Action Scheduler for reliable delivery. The subscriber must have a phone number on file and SMS opt-in enabled.')

# ── 7.3 Add Tag ──
add_heading2('7.3 Add Tag Step')

add_heading3('What It Does')
add_para(
    'The Add Tag step adds a specified tag to the subscriber\'s profile. Tags are useful for segmentation, '
    'tracking which flows a subscriber has been through, and triggering condition-based logic.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Add Tag from the step type dropdown.')
add_step(3, 'Enter the Tag name you want to add (e.g., "completed-welcome-series", "vip-customer", "cart-abandoner").')
add_step(4, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Tag — ', 'the tag string to add. If the subscriber already has this tag, no duplicate is created.')

add_heading3('Example Use Case')
add_para(
    'At the end of your welcome series flow, add the tag "completed-welcome-series". '
    'You can then use this tag in segments to exclude subscribers who have already received the welcome series, '
    'or use it in condition branches in other flows.'
)

# ── 7.4 Remove Tag ──
add_heading2('7.4 Remove Tag Step')

add_heading3('What It Does')
add_para(
    'The Remove Tag step removes a specified tag from the subscriber\'s profile. '
    'This is useful for resetting states, removing temporary labels, or updating subscriber categorisation as they progress through flows.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Remove Tag from the step type dropdown.')
add_step(3, 'Enter the Tag name you want to remove.')
add_step(4, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Tag — ', 'the tag string to remove. If the subscriber does not have this tag, the step is silently skipped.')

add_heading3('Example Use Case')
add_para(
    'In a post-purchase flow, remove the tag "cart-abandoner" from the subscriber since they have now completed a purchase. '
    'This keeps your tags clean and ensures the subscriber isn\'t included in abandonment-related segments.'
)

# ── 7.5 Update Field ──
add_heading2('7.5 Update Field Step')

add_heading3('What It Does')
add_para(
    'The Update Field step sets or changes a custom field value on the subscriber\'s profile. '
    'Custom fields store arbitrary data like preferences, scores, or status flags.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Update Field from the step type dropdown.')
add_step(3, 'Enter the Field Name (e.g., "loyalty_tier", "onboarding_status", "preferred_category").')
add_step(4, 'Enter the Field Value to set (e.g., "gold", "completed", "electronics").')
add_step(5, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Field Name — ', 'the name of the custom field to update. Both the field name and value are sanitised for safety.')
add_bold_bullet('Field Value — ', 'the value to set for the custom field. This overwrites any existing value for the field.')

add_heading3('Example Use Case')
add_para(
    'After a subscriber completes your welcome series, update the custom field "onboarding_status" to "completed". '
    'You can then use this field in segments and condition branches to personalise future communications.'
)

# ── 7.6 Condition Branch ──
add_heading2('7.6 Condition Branch Step')

add_heading3('What It Does')
add_para(
    'The Condition Branch step evaluates one or more rules against the subscriber\'s data and routes them '
    'down a "Yes" path or "No" path based on whether all the rules match. This enables you to create '
    'personalised, branching flows that adapt to each subscriber\'s behaviour and attributes.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Condition Branch from the step type dropdown.')
add_step(3, 'Define your rules. Each rule specifies a field, operator, and value. All rules must match for the "Yes" path (AND logic).')
add_step(4, 'Specify the Yes Step — the step to jump to if all rules match.')
add_step(5, 'Specify the No Step — the step to jump to if any rule does not match.')
add_step(6, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Rules — ', 'an array of conditions to evaluate. Each rule has a field, operator, and value.')
add_para('Available operators:')
add_bullet('equals / not_equals — exact match or mismatch')
add_bullet('greater_than / less_than — numeric comparisons')
add_bullet('contains / not_contains — substring matching')
add_bullet('is_blank / is_not_blank — empty value checks')
add_bullet('has_tag / not_has_tag — tag presence checks')
add_bullet('is_true / is_false — boolean checks')
add_bold_bullet('Yes Step ID — ', 'the step to advance to if all conditions match.')
add_bold_bullet('No Step ID — ', 'the step to advance to if any condition does not match.')

add_para('The condition evaluator can check:')
add_bullet('Direct subscriber fields (e.g., first_name, email, total_orders, total_spent)')
add_bullet('Custom fields (any custom field stored in the subscriber\'s profile)')
add_bullet('Tags (whether the subscriber has or doesn\'t have a specific tag)')
add_bullet('Engagement data (whether the subscriber has opened or clicked emails, checked via the ams_sends table)')

add_heading3('Example Use Case')
add_para(
    'In a welcome series, after the third email, add a condition branch that checks if the subscriber has '
    'placed at least 1 order (total_orders greater_than 0). If yes, send a "Thank you for your purchase" email. '
    'If no, send a "Here\'s 10% off your first order" discount email.'
)

add_tip('You can chain multiple condition branches for complex logic. For example, first check if they purchased, then check which product category they bought from to send targeted recommendations.')

# ── 7.7 Wait ──
add_heading2('7.7 Wait Step')

add_heading3('What It Does')
add_para(
    'The Wait step pauses the flow for a specified duration before advancing to the next step. '
    'The delay is processed by Action Scheduler, meaning it is reliable even if your site has intermittent traffic.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Wait from the step type dropdown.')
add_step(3, 'Enter the Delay Value (a number).')
add_step(4, 'Select the Delay Unit from the dropdown: minutes, hours, days, or weeks.')
add_step(5, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_bold_bullet('Delay Value — ', 'the numeric amount of time to wait.')
add_bold_bullet('Delay Unit — ', 'the time unit. Options: minutes, hours, days, weeks.')

add_para('Examples:')
add_bullet('1 hour = the next step runs 1 hour after the Wait step is reached')
add_bullet('2 days = the next step runs 2 days later')
add_bullet('1 week = the next step runs 7 days later')

add_heading3('Example Use Case')
add_para(
    'In a post-purchase flow, use a Wait step of 7 days between the thank-you email and the review request email. '
    'This gives the customer time to receive and try the product before asking for a review.'
)

add_tip('Use shorter waits (30 minutes to 1 hour) for time-sensitive flows like abandoned cart recovery, and longer waits (2-7 days) for nurturing sequences like welcome series.')

# ── 7.8 Exit Flow ──
add_heading2('7.8 Exit Flow Step')

add_heading3('What It Does')
add_para(
    'The Exit Flow step immediately ends the subscriber\'s enrolment in the flow. '
    'The enrolment is marked as exited with the reason "exit_step". '
    'No further steps in the flow will be executed for this subscriber.'
)

add_heading3('How to Add It to a Flow')
add_step(1, 'In the flow editor, click Add Step.')
add_step(2, 'Select Exit Flow from the step type dropdown.')
add_step(3, 'No additional configuration is needed.')
add_step(4, 'Click Save Flow to save your changes.')

add_heading3('Configuration Options')
add_para('This step has no configuration options. It simply ends the flow for the subscriber.')

add_heading3('Example Use Case')
add_para(
    'In a win-back flow, use a condition branch to check if the subscriber has made a purchase. '
    'If yes, route them to an Exit Flow step — they\'ve been won back and don\'t need further win-back emails. '
    'If no, continue with the next win-back message.'
)

add_tip('Exit Flow steps are especially useful at the end of Condition Branch "Yes" or "No" paths to prevent subscribers from accidentally continuing through unintended steps.')


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 8: PRE-BUILT FLOW TEMPLATES
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 8: Pre-Built Flow Templates')

add_heading2('8.1 What Are Flow Templates?')
add_para(
    'Flow templates are pre-configured automation flows that you can import with one click. '
    'They provide a ready-made starting point for common marketing automations, complete with triggers, steps, '
    'wait times, email content, and condition logic already configured. After importing a template, '
    'you can customise every aspect to match your brand and strategy.'
)

add_heading2('8.2 Available Templates')
add_para('Apotheca Marketing Suite includes 5 pre-built flow templates:')

add_heading3('8.2.1 Welcome Series')
add_bold_para('Template name: ', 'Welcome Series')
add_bold_para('Trigger: ', 'Welcome Series (fires on new subscriber confirmation)')
add_bold_para('Steps: ', '9 steps')
add_para('This template creates a comprehensive welcome sequence:')
add_bullet('Immediate: Welcome email greeting the new subscriber by name')
add_bullet('Wait 2 days: Brand story email sharing your company\'s mission and values')
add_bullet('Wait 3 days: Trending products email showcasing popular items')
add_bullet('Wait 2 days: Exclusive discount email with a 10% off code (WELCOME10)')
add_bullet('Wait 3 days: Condition branch checking if the subscriber has purchased')
add_bullet('If no purchase: Send SMS reminder (for subscribers with SMS opt-in)')

add_heading3('8.2.2 Abandoned Cart Recovery')
add_bold_para('Template name: ', 'Abandoned Cart Recovery')
add_bold_para('Trigger: ', 'Abandoned Cart')
add_bold_para('Steps: ', '5 steps')
add_para('This template creates a 3-email recovery sequence:')
add_bullet('Wait 1 hour: Cart reminder email with a direct link back to the cart')
add_bullet('Wait 23 hours: Social proof email featuring customer reviews and testimonials')
add_bullet('Wait 24 hours: Final discount email offering 10% off with code CART10')

add_heading3('8.2.3 Browse Abandonment')
add_bold_para('Template name: ', 'Browse Abandonment')
add_bold_para('Trigger: ', 'Browse Abandonment')
add_bold_para('Steps: ', '2 steps (with waits)')
add_para('This template creates a 2-email browsing follow-up:')
add_bullet('Wait 30 minutes: Product reminder email highlighting the items they viewed')
add_bullet('Wait 24 hours: Curated recommendations email with related products')

add_heading3('8.2.4 Post-Purchase Thank You')
add_bold_para('Template name: ', 'Post-Purchase Thank You')
add_bold_para('Trigger: ', 'Post-Purchase')
add_bold_para('Steps: ', '2 steps (with waits)')
add_para('This template creates a post-purchase follow-up:')
add_bullet('Immediate: Thank-you email confirming the order and expressing gratitude')
add_bullet('Wait 7 days: Review request email with a star rating widget (1–5 stars clickable)')

add_heading3('8.2.5 Win-Back Campaign')
add_bold_para('Template name: ', 'Win-Back Campaign')
add_bold_para('Trigger: ', 'Win-Back (90 days since last order)')
add_bold_para('Steps: ', '3 steps (with waits)')
add_para('This template creates a 3-touch re-engagement sequence:')
add_bullet('Day 0: "We miss you" email reconnecting with the inactive customer')
add_bullet('Wait 15 days: Discount email offering 15% off with code COMEBACK15')
add_bullet('Wait 15 days: SMS reminder as a final touchpoint')

add_heading2('8.3 How to Import a Template')
add_step(1, 'Navigate to Flows in the Apotheca Marketing menu.')
add_step(2, 'On the flow list page, look for the Templates section or the Import Template button.')
add_step(3, 'Browse the available templates and click Import on the template you want to use.')
add_step(4, 'The flow builder will create a new flow pre-populated with the template\'s trigger, steps, and configuration.')
add_step(5, 'Review the imported flow and make any customisations needed (see section 8.4 below).')
add_step(6, 'Click Save Flow to save the imported template.')
add_step(7, 'Set the flow status to Active when you\'re ready to go live.')

add_heading2('8.4 How to Customise Imported Templates')
add_para('After importing a template, you should review and customise it to match your brand:')

add_bold_bullet('Email content: ', 'Replace the placeholder email body text with your own brand voice, product images, and calls to action. Update subject lines to match your style.')
add_bold_bullet('Discount codes: ', 'Templates include example discount codes (e.g., WELCOME10, CART10, COMEBACK15). Create actual WooCommerce coupon codes matching these, or change the codes in the email content to match your existing coupons.')
add_bold_bullet('Wait times: ', 'Adjust the wait durations between steps to match your preferred cadence. For example, you might want shorter gaps for abandoned cart recovery or longer gaps for welcome series.')
add_bold_bullet('Personalisation tokens: ', 'Templates already use {{first_name}}, {{site_name}}, and {{site_url}} tokens. You can add more tokens like {{total_orders}} or {{rfm_segment}} where appropriate.')
add_bold_bullet('Condition logic: ', 'Review any condition branches and ensure the rules match your business logic. For example, the welcome series template checks if a subscriber has purchased — verify the condition operators are correct for your needs.')
add_bold_bullet('SMS steps: ', 'Some templates include SMS steps. If you don\'t use SMS marketing, remove these steps or replace them with email alternatives.')

add_tip('Import a template, customise it for your brand, and then save it as a starting point. You can always create additional flows from scratch if the templates don\'t match your needs.')

add_important('Templates include example email content that is generic. Always replace template email body text with your own branded content before activating the flow.')


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 9: SMART SEGMENTATION
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 9: Smart Segmentation')

add_heading2('9.1 What Are Segments?')
add_para(
    'Segments are dynamic groups of subscribers defined by rules and conditions. Unlike static lists, '
    'segments automatically update as subscriber data changes — subscribers are included or excluded based on '
    'whether they match the segment\'s conditions at the time of evaluation. This means your segments are always current.'
)
add_para(
    'Segments are the backbone of targeted marketing. You can use segments to:'
)
add_bullet('Target specific audiences in email campaigns')
add_bullet('Filter the subscriber list to find particular groups')
add_bullet('Combine with flows and triggers for hyper-targeted automations')
add_bullet('Analyse customer behaviour and engagement patterns')

add_heading2('9.2 Navigating to the Segments Page')
add_step(1, 'In your WordPress admin dashboard, look for the Apotheca Marketing menu in the left sidebar.')
add_step(2, 'Click Segments to open the Segments management page.')
add_step(3, 'You will see a table listing all existing segments with their name, subscriber count, and the last time the count was recalculated.')

add_heading2('9.3 Creating a New Segment')
add_step(1, 'On the Segments page, click the Create New Segment button.')
add_step(2, 'Enter a descriptive Segment Name (e.g., "High-Value VIP Customers", "Inactive 90+ Days", "Email Openers Last 30 Days").')
add_step(3, 'Build your conditions using the visual condition builder (see sections 9.4 and 9.5 below).')
add_step(4, 'Click the Preview Count button to see how many subscribers currently match your conditions.')
add_step(5, 'Review the count. If it looks right, click Save Segment.')
add_step(6, 'Your segment is now live and will be automatically recalculated every 6 hours.')

add_heading2('9.4 The 25+ Condition Types')
add_para(
    'Apotheca Marketing Suite supports over 25 condition types, organised into three categories. '
    'Each condition type has specific operators available to it.'
)

add_heading3('9.4.1 Subscriber Data Conditions')
add_para('These conditions evaluate core subscriber profile data:')
add_bold_bullet('Email Domain — ', 'filter by the domain portion of the email address. Operators: is, is_not, contains. Example: "email domain is gmail.com" to find Gmail users.')
add_bold_bullet('First Name — ', 'filter by the subscriber\'s first name. Operators: is, is_not, is_blank, is_not_blank, contains.')
add_bold_bullet('Tag — ', 'check whether the subscriber has or does not have specific tags. Operators: has, does_not_have. You can specify multiple tags separated by commas.')
add_bold_bullet('Custom Field — ', 'evaluate any custom field stored on the subscriber profile. Requires specifying the field name. Operators: equals, not_equals, contains, greater_than, less_than.')
add_bold_bullet('Source — ', 'filter by how the subscriber was acquired (e.g., "checkout", "registration", "import"). Operators: is, is_not.')
add_bold_bullet('GDPR Consent — ', 'check whether the subscriber has given GDPR consent. Operators: is_true, is_false.')
add_bold_bullet('Subscribed Date — ', 'filter by when the subscriber signed up. Operators: before, after, within_last_X_days, more_than_X_days_ago.')
add_bold_bullet('Predicted CLV — ', 'filter by the subscriber\'s predicted Customer Lifetime Value. Operators: greater_than, less_than, equals, between.')
add_bold_bullet('Churn Risk Score — ', 'filter by the subscriber\'s churn risk score (0–100). Operators: greater_than, less_than, equals.')
add_bold_bullet('RFM Segment — ', 'filter by the subscriber\'s named RFM segment. Operators: is, is_not. Values: Champions, Big Spenders, Loyal, New Customers, Potential, At Risk, About to Sleep, Lost.')

add_heading3('9.4.2 E-Commerce Data Conditions')
add_para('These conditions evaluate WooCommerce purchase data:')
add_bold_bullet('Total Orders — ', 'the subscriber\'s total number of orders. Operators: greater_than, less_than, equals, between.')
add_bold_bullet('Total Spent — ', 'the subscriber\'s total amount spent across all orders. Operators: greater_than, less_than, equals, between.')
add_bold_bullet('Average Order Value — ', 'calculated as total spent divided by total orders. Operators: greater_than, less_than, equals, between.')
add_bold_bullet('Last Order Date — ', 'the date of the subscriber\'s most recent order. Operators: before, after, within_last_X_days, more_than_X_days_ago.')
add_bold_bullet('Purchased Product — ', 'whether the subscriber has purchased a specific product (by ID or SKU). Operators: has, has_not.')
add_bold_bullet('Purchased Category — ', 'whether the subscriber has purchased from a specific WooCommerce product category (by slug). Operators: has, has_not.')
add_bold_bullet('Last Order Status — ', 'the status of the subscriber\'s most recent order. Operators: is, is_not. Values: completed, processing, on-hold, refunded, cancelled, failed, etc.')
add_bold_bullet('Used Coupon — ', 'whether the subscriber has ever used a specific coupon code. Operators: has, has_not.')

add_heading3('9.4.3 Engagement Data Conditions')
add_para('These conditions evaluate email and SMS engagement:')
add_bold_bullet('Opened Campaign — ', 'whether the subscriber opened a specific email campaign (by campaign ID). Operators: has, has_not.')
add_bold_bullet('Clicked Campaign — ', 'whether the subscriber clicked a link in a specific email campaign. Operators: has, has_not.')
add_bold_bullet('Opened Any Email — ', 'whether the subscriber has ever opened any email, or opened within a time period. Operators: ever, never, within_last_X_days.')
add_bold_bullet('Clicked Any Email — ', 'whether the subscriber has ever clicked any email link, or clicked within a time period. Operators: ever, never, within_last_X_days.')
add_bold_bullet('SMS Opt-In — ', 'whether the subscriber has a phone number (indicating SMS opt-in). Operators: is_true, is_false.')
add_bold_bullet('Email Bounce Status — ', 'the bounce status of the subscriber\'s most recent email. Operators: is, is_not. Values: none, soft, hard.')

add_heading2('9.5 AND/OR Logic and Nested Groups')
add_para(
    'The condition builder supports nested groups with AND/OR logic, allowing you to create sophisticated '
    'segments with complex rules. You can nest groups up to 3 levels deep.'
)

add_heading3('How AND/OR Logic Works')
add_bold_bullet('AND logic — ', 'ALL conditions in the group must match. A subscriber is included only if every rule evaluates to true. Example: "total_orders greater_than 5 AND total_spent greater_than 500" matches only subscribers with both high order count AND high spending.')
add_bold_bullet('OR logic — ', 'ANY condition in the group can match. A subscriber is included if at least one rule evaluates to true. Example: "source is checkout OR source is registration" matches subscribers who came from either channel.')

add_heading3('How to Create Nested Groups')
add_step(1, 'In the condition builder, your top-level group defaults to AND logic.')
add_step(2, 'Click the logic toggle (AND/OR) to switch between AND and OR logic for the group.')
add_step(3, 'Click Add Condition to add a new rule to the current group.')
add_step(4, 'Click Add Group to create a nested sub-group within the current group.')
add_step(5, 'Configure the nested group\'s logic (AND/OR) and add conditions to it.')
add_step(6, 'You can nest groups up to 3 levels deep for complex combinations.')

add_para('Example of a nested segment: "VIP Customers Who Are Engaged":')
add_bullet('Group 1 (AND):')
add_bullet('  — total_spent greater_than 1000')
add_bullet('  — Group 2 (OR):')
add_bullet('      — opened_any_email within_last_X_days 30')
add_bullet('      — clicked_any_email within_last_X_days 30')
add_para('This segment finds subscribers who have spent more than $1,000 AND have either opened or clicked an email in the last 30 days.')

add_heading2('9.6 Live Subscriber Count Preview')
add_para(
    'The segment builder includes a Preview Count button that lets you test your segment conditions '
    'in real time before saving.'
)
add_step(1, 'Build or modify your segment conditions.')
add_step(2, 'Click the Preview Count button.')
add_step(3, 'The system evaluates all subscribers against your conditions and displays the matching count.')
add_step(4, 'Use this to verify your segment targets the right audience before saving.')

add_tip('If the count seems too high or too low, review your conditions and logic. A common mistake is using OR when you mean AND (which broadens the segment) or AND when you mean OR (which narrows it).')

add_heading2('9.7 Automatic Recalculation')
add_para(
    'Segments are automatically recalculated every 6 hours via an Action Scheduler background job. '
    'This ensures your segment counts stay current as subscriber data changes throughout the day.'
)
add_para('Here is how the recalculation process works:')
add_bullet('Every 6 hours, the segment calculator runs in the background.')
add_bullet('It processes subscribers in batches of 500 for memory efficiency.')
add_bullet('Each subscriber is evaluated against every segment\'s conditions using the ConditionEvaluator.')
add_bullet('The subscriber count for each segment is updated and cached.')
add_bullet('The "Last Calculated" timestamp is updated on the segments list page.')

add_para('You can also trigger an on-demand count by clicking the Preview Count button when editing a segment.')

add_heading2('9.8 Editing and Deleting Segments')

add_heading3('Editing a Segment')
add_step(1, 'Navigate to the Segments page.')
add_step(2, 'Click the name of the segment you want to edit, or click the Edit button.')
add_step(3, 'Modify the segment name, conditions, or logic as needed.')
add_step(4, 'Click Preview Count to verify the updated conditions.')
add_step(5, 'Click Save Segment to apply your changes.')

add_heading3('Deleting a Segment')
add_step(1, 'Navigate to the Segments page.')
add_step(2, 'Click the Delete button next to the segment you want to remove.')
add_step(3, 'Confirm the deletion when prompted.')

add_important('Deleting a segment is permanent and cannot be undone. If any flows or campaigns reference this segment, they will no longer be able to use it. Review your flows and campaigns before deleting a segment.')


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 10: RFM SCORING
# ════════════════════════════════════════════════════════════════════════════

add_heading1('Chapter 10: RFM Scoring')

add_heading2('10.1 What Is RFM Scoring?')
add_para(
    'RFM scoring is a customer segmentation technique based on three metrics that predict future buying behaviour:'
)
add_bold_bullet('Recency (R) — ', 'How recently did the customer place an order? Customers who purchased recently are more likely to buy again.')
add_bold_bullet('Frequency (F) — ', 'How often does the customer order? Frequent buyers are more loyal and valuable.')
add_bold_bullet('Monetary (M) — ', 'How much has the customer spent in total? High spenders represent your most valuable customers.')

add_para(
    'Each metric is scored on a scale of 1 to 5, where 5 is the best. '
    'A subscriber with an RFM score of "555" is your best customer (bought recently, buys often, spends a lot), '
    'while a score of "111" indicates a lost customer (hasn\'t bought in ages, rarely buys, low spending).'
)

add_heading2('10.2 How Scores Are Calculated')
add_para(
    'The RFM scoring engine runs nightly via Action Scheduler, processing all subscribers who have at least one order. '
    'Subscribers are processed in batches of 500 for memory efficiency.'
)

add_heading3('Recency Score')
add_para('Recency is scored based on the number of days since the subscriber\'s last order:')
add_bullet('Score 5: 0–14 days since last order (most recent)')
add_bullet('Score 4: 15–30 days since last order')
add_bullet('Score 3: 31–60 days since last order')
add_bullet('Score 2: 61–180 days since last order')
add_bullet('Score 1: 180+ days since last order (least recent)')

add_heading3('Frequency Score')
add_para(
    'Frequency is scored using quintile ranking. The system collects all subscribers\' total order counts, '
    'sorts them, and divides them into 5 equal groups (quintiles). Your score depends on which quintile you fall into relative to all other subscribers:'
)
add_bullet('Score 5: Top 20% — highest order count')
add_bullet('Score 4: 60th–80th percentile')
add_bullet('Score 3: 40th–60th percentile')
add_bullet('Score 2: 20th–40th percentile')
add_bullet('Score 1: Bottom 20% — lowest order count')

add_heading3('Monetary Score')
add_para(
    'Monetary is also scored using quintile ranking, using the same approach as Frequency but based on total amount spent:'
)
add_bullet('Score 5: Top 20% — highest total spending')
add_bullet('Score 4: 60th–80th percentile')
add_bullet('Score 3: 40th–60th percentile')
add_bullet('Score 2: 20th–40th percentile')
add_bullet('Score 1: Bottom 20% — lowest total spending')

add_heading2('10.3 The 8 Named RFM Segments')
add_para(
    'Based on the combination of R, F, and M scores, each subscriber is assigned to one of 8 named segments. '
    'Segments are evaluated in priority order — the first matching rule wins.'
)

# Champions
add_heading3('1. Champions')
add_bold_para('Criteria: ', 'R ≥ 4, F ≥ 4, M ≥ 4')
add_para(
    'Your best customers. They bought recently, buy frequently, and spend the most. These are your brand advocates — '
    'nurture them with exclusive offers, early access to new products, and VIP treatment. They are the most likely to '
    'respond to upsells and cross-sells.'
)

# Big Spenders
add_heading3('2. Big Spenders')
add_bold_para('Criteria: ', 'M = 5 (any R and F)')
add_para(
    'Customers in the top 20% for spending, regardless of how recently or frequently they buy. '
    'They may not buy often, but when they do, they spend big. Target them with premium products and high-ticket offers.'
)

# Loyal
add_heading3('3. Loyal')
add_bold_para('Criteria: ', 'R ≥ 3, F ≥ 4')
add_para(
    'Frequent buyers who have purchased recently. They may not be the highest spenders, but they are consistent. '
    'Reward their loyalty with loyalty programmes, referral incentives, and personalised recommendations.'
)

# New Customers
add_heading3('4. New Customers')
add_bold_para('Criteria: ', 'R = 5, F = 1')
add_para(
    'Customers who placed their very first order recently. This is a critical segment — their second purchase '
    'determines whether they become loyal or lapse. Nurture them with welcome offers, product education, '
    'and a seamless post-purchase experience.'
)

# Potential
add_heading3('5. Potential')
add_bold_para('Criteria: ', 'R ≥ 4, F ≤ 2')
add_para(
    'Recent buyers who haven\'t purchased frequently yet. They show promise but need encouragement. '
    'Target them with incentives to make their next purchase, product bundles, and "customers also bought" recommendations.'
)

# At Risk
add_heading3('6. At Risk')
add_bold_para('Criteria: ', 'R ≤ 2, F ≥ 3, M ≥ 3')
add_para(
    'Previously valuable customers who haven\'t purchased recently. These were once frequent, high-value buyers '
    'but are slipping away. This is an urgent segment — reach out with personalised win-back offers, '
    'ask for feedback, and remind them of what they\'re missing.'
)

# About to Sleep
add_heading3('7. About to Sleep')
add_bold_para('Criteria: ', 'R ≤ 3, F ≤ 2, M ≤ 2')
add_para(
    'Low-activity customers showing early signs of lapsing. They haven\'t bought recently, don\'t buy often, '
    'and don\'t spend much. Try re-engaging them with special offers or ask them to update their preferences. '
    'If they don\'t respond, they may transition to Lost.'
)

# Lost
add_heading3('8. Lost')
add_bold_para('Criteria: ', 'R = 1, F ≤ 2')
add_para(
    'Customers who haven\'t purchased in over 180 days and never bought frequently. '
    'These are the hardest to win back. Consider a last-resort aggressive offer, or accept that they may have churned '
    'and focus your resources on higher-potential segments.'
)

add_tip('There is also a fallback "Other" segment for subscribers whose scores don\'t match any of the 8 named segments. This is rare but can occur with unusual score combinations.')

add_heading2('10.4 Using RFM Segments in Flows and Segments')

add_heading3('In Flows')
add_para('You can use RFM data in flows in two ways:')
add_bold_bullet('RFM Change trigger — ', 'automatically start a flow when a subscriber\'s RFM segment changes (see Chapter 6, Section 6.7). This lets you react in real time to customer behaviour shifts.')
add_bold_bullet('Condition Branch step — ', 'use a condition branch within a flow to check the subscriber\'s RFM segment and route them down different paths. For example, send Champions a premium offer and Potential customers a starter discount.')

add_heading3('In Segments')
add_para('You can filter by RFM data in the segment builder:')
add_bold_bullet('RFM Segment condition — ', 'use the "rfm_segment" condition type with the "is" or "is_not" operator to include or exclude subscribers in a specific RFM segment. Example: "rfm_segment is Champions" to create a segment of only your best customers.')
add_bold_bullet('Predicted CLV condition — ', 'filter by predicted customer lifetime value (see Section 10.6). Example: "predicted_clv greater_than 500" to find high-value customers.')
add_bold_bullet('Churn Risk Score condition — ', 'filter by churn risk score. Example: "churn_risk_score greater_than 70" to find customers at high risk of churning.')

add_heading2('10.5 Viewing RFM Data on Subscriber Profiles')
add_para(
    'Each subscriber\'s profile stores their RFM data, which is updated nightly. On a subscriber\'s profile, you can view:'
)
add_bold_bullet('RFM Score — ', 'the 3-digit score (e.g., "543") representing Recency, Frequency, and Monetary scores.')
add_bold_bullet('RFM Segment — ', 'the named segment the subscriber belongs to (e.g., "Champions", "At Risk", "Lost").')

add_para(
    'This data is recalculated automatically every night. You do not need to manually trigger recalculation.'
)

add_heading2('10.6 Predictive CLV and Churn Risk Scores')
add_para(
    'In addition to RFM scoring, Apotheca Marketing Suite includes a predictive analytics engine that calculates '
    'two additional metrics for each subscriber with order history. These are also recalculated nightly.'
)

add_heading3('Predicted Customer Lifetime Value (CLV)')
add_para(
    'The predicted CLV estimates how much revenue a subscriber is expected to generate over the next 12 months, based on their historical ordering pattern.'
)
add_para('The formula is:')
add_bold_para('Predicted CLV = ', 'Average Order Value × (365 ÷ Average Gap Between Orders)')
add_para(
    'For example, if a subscriber spends an average of $50 per order and orders roughly every 30 days, '
    'their predicted 12-month CLV would be $50 × (365 ÷ 30) = $608.33.'
)

add_heading3('Churn Risk Score (0–100)')
add_para(
    'The churn risk score estimates the likelihood that a subscriber will not purchase again. '
    'It ranges from 0 (low risk) to 100 (high risk).'
)
add_para('The formula is:')
add_bold_para('Churn Risk = ', '(Days Since Last Order ÷ Average Gap Between Orders) × 50, capped at 100')
add_para(
    'For example, if a subscriber\'s average order gap is 30 days and they last ordered 60 days ago, '
    'their churn risk is (60 ÷ 30) × 50 = 100 (high risk). If they last ordered 15 days ago, '
    'the churn risk is (15 ÷ 30) × 50 = 25 (low risk).'
)
add_para(
    'For subscribers with only one order, a default average gap of 90 days is used in the calculation.'
)

add_heading3('Predicted Next Order Date')
add_para(
    'The predictive engine also estimates when each subscriber is likely to place their next order, '
    'based on their last order date plus their average gap between orders. '
    'If the predicted date is already in the past, it is projected forward to the next expected date.'
)

add_important('A churn risk score above 70 indicates a high-risk subscriber. Consider targeting these subscribers with re-engagement campaigns using the churn_risk_score condition in segments or the RFM Change trigger in flows.')

add_tip('Combine RFM segments with predictive CLV for powerful targeting. For example, create a segment of "Potential" customers with a predicted CLV above $500 — these are your best growth opportunities.')


# ── Save ────────────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Apotheca_User_Manual_Part2.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
