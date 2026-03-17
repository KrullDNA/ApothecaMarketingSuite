#!/usr/bin/env python3
"""Generate Apotheca User Manual Part 3 as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(13)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text):
    doc.add_paragraph(text)


def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)
    p.add_run(text)


def add_important(text):
    p = doc.add_paragraph()
    run = p.add_run('IMPORTANT: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run(text)


def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {number}. ')
    run.bold = True
    p.add_run(text)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(text)


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Apotheca Marketing Suite')
run.font.size = Pt(28)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('User Manual \u2014 Part 3')
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chapters 11\u201315: Forms, SMS, Email Editor, and Reviews')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0.0')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading('Table of Contents', 1)
toc_items = [
    ('Chapter 11', 'Pop-Up & Opt-In Forms'),
    ('Chapter 12', 'Spin-to-Win Forms'),
    ('Chapter 13', 'SMS Marketing'),
    ('Chapter 14', 'Visual Email Editor'),
    ('Chapter 15', 'Reviews Integration'),
]
for ch, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{ch}: ')
    run.bold = True
    p.add_run(title)

doc.add_page_break()

# ============================================================
# CHAPTER 11: POP-UP & OPT-IN FORMS
# ============================================================
add_heading('Chapter 11: Pop-Up & Opt-In Forms', 1)

add_para(
    'Apotheca Marketing Suite includes a powerful form builder that lets you create pop-up and opt-in '
    'forms to capture email addresses and phone numbers from your website visitors. You can choose from '
    'six different form types, configure targeting rules to control who sees each form, and set up '
    'success actions like tagging subscribers or enrolling them in automated flows.'
)

# -- 11.1 The Six Form Types --
add_heading('The Six Form Types', 2)

add_para(
    'Each form type is designed for a different use case. Choose the type that best matches where and '
    'how you want to capture subscribers.'
)

add_heading('Modal Pop-up', 3)
add_para(
    'A modal is a centred overlay that appears on top of the page content with a dimmed background. '
    'It demands attention and is the most common form type for capturing email subscribers.'
)
add_bullet('Best for: Homepage welcome offers, site-wide promotions, exit-intent captures.')
add_bullet('The visitor must close the modal or submit the form to continue browsing.')

add_heading('Flyout / Slide-in', 3)
add_para(
    'A flyout slides in from the bottom-left or bottom-right corner of the screen. It is less intrusive '
    'than a modal because it does not cover the full page.'
)
add_bullet('Best for: Gentle nudges on blog posts, content pages, or product pages.')
add_bullet('You can choose the flyout position (bottom-left or bottom-right) in the Design step.')

add_heading('Embedded', 3)
add_para(
    'An embedded form is placed directly inside your page content using a shortcode. It does not pop up '
    'or slide in \u2014 it sits inline with the rest of the page.'
)
add_bullet('Best for: Newsletter sign-ups in blog sidebars, footer areas, or dedicated landing pages.')
add_bullet('Use the shortcode [ams_form id="123"] where 123 is the form ID.')

add_heading('Full Page', 3)
add_para(
    'A full-page form takes over the entire browser window with a full-screen overlay. This is the most '
    'aggressive form type and should be used sparingly.'
)
add_bullet('Best for: Major promotions, product launches, or high-value lead magnets.')
add_bullet('Visitors must interact with the form or close it to continue.')

add_heading('Sticky Bar', 3)
add_para(
    'A sticky bar is a thin banner that appears at the top or bottom of the browser window and stays '
    'fixed as the visitor scrolls.'
)
add_bullet('Best for: Persistent, low-distraction announcements like free shipping thresholds or limited-time offers.')
add_bullet('You can choose the bar position (top or bottom) in the Design step.')

add_heading('Spin-to-Win', 3)
add_para(
    'A gamified modal that displays a prize wheel. The visitor enters their email, spins the wheel, and '
    'receives a coupon code if they win. This form type has additional configuration for prize segments '
    '(covered in detail in Chapter 12).'
)
add_bullet('Best for: Engaging first-time visitors, boosting conversion rates with gamification.')
add_bullet('Automatically generates WooCommerce coupon codes for winners.')

# -- 11.2 Creating a New Form --
add_heading('Creating a New Form', 2)

add_step(1, 'In the WordPress admin menu, go to Apotheca Marketing > Forms.')
add_step(2, 'Click the + New Form button in the top-right corner.')
add_step(3, 'Enter a Form Name (e.g., "Homepage Pop-up" or "Blog Sidebar Signup").')
add_step(4, 'Select a Form Type from the dropdown (Modal Pop-up, Flyout, Embedded, Full Page, Sticky Bar, or Spin-to-Win).')
add_step(5, 'You will see a multi-step wizard with tabs: Fields, Design, Targeting, and Success. If you selected Spin-to-Win, an extra Spin Config tab appears between Fields and Design.')
add_step(6, 'Work through each tab to configure your form (detailed below).')
add_step(7, 'Click Save Form when you are finished.')

add_tip(
    'New forms are saved in Draft status. You must go back to the form list and click Activate '
    'to make the form live on your site.'
)

# -- 11.3 Form Builder Walkthrough --
add_heading('Form Builder Interface Walkthrough', 2)

add_para(
    'The form builder uses a multi-step wizard with a live preview panel on the right side. '
    'You can toggle between Desktop and Mobile preview to see how the form looks on different screen sizes.'
)

add_heading('Step 1: Fields', 3)
add_para(
    'Configure which fields appear on your form. The Email Address field is always required and cannot be removed. '
    'You can add, remove, and reorder additional fields.'
)

add_para('The nine available field types are:')
add_bullet('Email Address \u2014 Required. Captures the subscriber\'s email. Always present.', 'Email Address:')
add_bullet('Collects the subscriber\'s phone or mobile number for SMS marketing.', 'Phone / Mobile:')
add_bullet('Captures the subscriber\'s first name for personalisation.', 'First Name:')
add_bullet('Captures the subscriber\'s last name.', 'Last Name:')
add_bullet('Displays two dropdowns (Month and Day) for birthday marketing campaigns.', 'Birthday:')
add_bullet('Presents multiple options where the visitor selects one. Enter options as a comma-separated list.', 'Radio Buttons:')
add_bullet('Presents multiple options where the visitor can select more than one. Enter options as a comma-separated list.', 'Checkboxes:')
add_bullet('Presents a dropdown menu of options. Enter options as a comma-separated list.', 'Dropdown:')
add_bullet('A field that is invisible to the visitor. Useful for tracking source or campaign data.', 'Hidden Field:')

add_para('For each field you can set:')
add_bullet('A display label (e.g., "Your Email").')
add_bullet('Placeholder text shown inside the input before the visitor types.')
add_bullet('Whether the field is required (the email field is always required).')
add_bullet('For radio, checkbox, and dropdown fields: a comma-separated list of options.')
add_bullet('For hidden fields: a fixed value.')

add_para(
    'Use the up and down arrow buttons on each field row to reorder fields. Click the X button to remove a field. '
    'Click + Add Field to add a new field at the bottom.'
)

add_heading('Step 2: Design', 3)
add_para('Customise the visual appearance of your form:')
add_bullet('Five colour settings: Background, Text, Button, Button Text, and Button Hover.')
add_bullet('Border Radius (px) \u2014 Controls the roundness of the form corners.')
add_bullet('Font Family \u2014 Choose from System Default, Montserrat, Open Sans, Roboto, Lato, Poppins, Playfair Display, Merriweather, Raleway, Nunito, or Inter.')
add_bullet('Title \u2014 The headline displayed on the form (e.g., "Subscribe to our newsletter").')
add_bullet('Description \u2014 Supporting text below the title.')
add_bullet('Button Text \u2014 The text on the submit button (e.g., "Subscribe" or "Get My Discount").')
add_bullet('Consent Text \u2014 A GDPR consent message displayed with a checkbox. HTML links are allowed (e.g., I agree to the Privacy Policy).')
add_bullet('Header Image URL \u2014 An optional image displayed at the top of the form.')
add_bullet('For Flyout forms: Flyout Position (Bottom Right or Bottom Left).')
add_bullet('For Sticky Bar forms: Bar Position (Top or Bottom).')

add_heading('Step 3: Targeting', 3)
add_para(
    'Targeting rules control who sees the form and when. The plugin evaluates targeting rules on every '
    'page load. Rules are split between server-side checks (page, device, segment) and client-side '
    'triggers (scroll depth, time on page, exit intent).'
)

add_para('Available targeting rules:')

add_bold_para('Page Targeting')
add_bullet('All Pages \u2014 Show the form on every page of your site.', 'All Pages:')
add_bullet('Enter a comma-separated list of WordPress page IDs to limit the form to specific pages.', 'Specific Pages:')
add_bullet('Enter a comma-separated list of page IDs to hide the form on those pages.', 'Exclude Pages:')

add_bold_para('Device Targeting')
add_bullet('Show on all devices.', 'All Devices:')
add_bullet('Show only on desktop browsers (tablets and phones excluded).', 'Desktop Only:')
add_bullet('Show only on mobile devices and tablets.', 'Mobile Only:')

add_bold_para('Visitor Type')
add_bullet('Show to everyone.', 'All Visitors:')
add_bullet('Only show to first-time visitors (no cookie detected).', 'New Visitors Only:')
add_bullet('Only show to visitors who have been to your site before.', 'Returning Visitors Only:')

add_bold_para('Minimum Cart Value')
add_para(
    'Enter a dollar amount. The form will only appear if the visitor\'s WooCommerce cart total is '
    'at or above this value. Leave blank or set to 0 to ignore cart value.'
)

add_bold_para('Segment Match')
add_para(
    'Enter the ID of a subscriber segment. The form will only appear to visitors who are identified '
    'as subscribers and belong to that segment. Leave blank to show to all visitors regardless of segment.'
)

add_bold_para('Frequency Cap (days)')
add_para(
    'After a visitor sees or dismisses the form, it will not appear again for this many days. '
    'The default is 7 days. This is tracked via a browser cookie.'
)

add_bold_para('Display Triggers')
add_para('These triggers are evaluated in the visitor\'s browser:')
add_bullet('The form appears after the visitor scrolls past a certain percentage of the page. Enter a value from 1 to 100.', 'Scroll Depth (%):')
add_bullet('The form appears after the visitor has been on the page for the specified number of seconds.', 'Time on Page (seconds):')
add_bullet('The form appears when the visitor moves their mouse toward the browser\'s close/back buttons. Desktop only \u2014 does not work on mobile.', 'Exit Intent:')

add_tip(
    'You can combine multiple triggers. For example, set Scroll Depth to 50% and Time on Page to 5 seconds. '
    'The form will appear when the first trigger condition is met.'
)

add_heading('Step 4: Success', 3)
add_para('Configure what happens after a visitor submits the form:')

add_bold_para('On Submit Action')
add_bullet('Displays a message inside the form (e.g., "Thank you for subscribing!"). You customise the text.', 'Show Success Message:')
add_bullet('Sends the visitor to a URL of your choice (e.g., a thank-you page or a special offer page).', 'Redirect to URL:')

add_bold_para('Add Tags')
add_para(
    'Enter a comma-separated list of tags to apply to the subscriber (e.g., "newsletter, homepage-optin"). '
    'Tags are useful for segmentation and flow targeting.'
)

add_bold_para('Enrol in Flow')
add_para(
    'Enter the ID of an automated flow. The subscriber will be immediately enrolled in that flow upon submission. '
    'For example, you could enrol them in a Welcome Series flow.'
)

add_bold_para('Require Double Opt-In')
add_para(
    'When checked, the subscriber\'s status is set to Pending instead of Subscribed, and a double opt-in '
    'confirmation email is sent. The subscriber must click the confirmation link before they receive any '
    'marketing emails. This can also be controlled globally in Settings > GDPR.'
)

# -- 11.4 GDPR Consent --
add_heading('GDPR Consent Checkbox', 2)
add_para(
    'If you enter text in the Consent Text field in the Design step, a checkbox is displayed on the form. '
    'When the visitor checks it and submits, the plugin records their GDPR consent with a timestamp. '
    'You can use HTML in the consent text to link to your privacy policy.'
)
add_para('Example: I agree to the <a href="/privacy-policy">Privacy Policy</a> and consent to receiving marketing emails.')
add_important(
    'If you operate in the EU or sell to EU customers, you should enable either the consent checkbox '
    'or double opt-in (or both) to comply with GDPR requirements.'
)

# -- 11.5 Rate Limiting --
add_heading('Rate Limiting', 2)
add_para(
    'To prevent abuse, the plugin limits form submissions to 10 per IP address per minute. If a visitor '
    'exceeds this limit, they see the message "Too many submissions. Please try again later." This protects '
    'against spam bots and brute-force form attacks.'
)

# -- 11.6 Managing Forms --
add_heading('Managing Your Forms', 2)
add_para('The Forms list page shows all your forms with the following information:')
add_bullet('Name \u2014 The form name you entered.')
add_bullet('Type \u2014 The form type (Modal, Flyout, etc.).')
add_bullet('Status \u2014 Draft or Active.')
add_bullet('Views \u2014 How many times the form has been displayed to visitors.')
add_bullet('Submissions \u2014 How many times the form has been submitted.')
add_bullet('Rate \u2014 The conversion rate (submissions divided by views).')

add_para('From the Actions column you can:')
add_bullet('Edit \u2014 Open the form in the builder to make changes.')
add_bullet('Activate / Deactivate \u2014 Toggle the form between Active and Draft status.')
add_bullet('Delete \u2014 Permanently remove the form (you will be asked to confirm).')

add_tip(
    'Only forms with Active status are displayed on your website. Draft forms are saved but invisible to visitors.'
)

# -- 11.7 Embedding --
add_heading('Embedding Forms with Shortcodes', 2)
add_para(
    'Embedded forms are placed in your page content using a shortcode. To embed a form:'
)
add_step(1, 'Create a form with the Embedded type.')
add_step(2, 'Note the form\'s ID number (visible in the form list or in the URL when editing).')
add_step(3, 'In any WordPress page, post, or widget, insert the shortcode: [ams_form id="123"] (replace 123 with your form\'s ID).')
add_step(4, 'Make sure the form\'s status is Active.')

add_para('The embedded form will render inline where you placed the shortcode.')

doc.add_page_break()

# ============================================================
# CHAPTER 12: SPIN-TO-WIN FORMS
# ============================================================
add_heading('Chapter 12: Spin-to-Win Forms', 1)

add_para(
    'Spin-to-Win is a gamified form type that displays a prize wheel to your visitors. When a visitor enters '
    'their email address and submits the form, the wheel spins and lands on a randomly selected prize segment. '
    'If the prize includes a discount, the plugin automatically generates a unique WooCommerce coupon code '
    'and displays it to the visitor.'
)

add_heading('How Spin-to-Win Works', 2)
add_step(1, 'The visitor sees a modal pop-up with a wheel graphic and an email input field.')
add_step(2, 'The visitor enters their email address and clicks the submit button.')
add_step(3, 'The plugin creates or updates the subscriber record (just like any other form).')
add_step(4, 'The server determines the winning prize segment based on probability weights.')
add_step(5, 'If the prize has a discount value, a unique WooCommerce coupon is generated.')
add_step(6, 'The wheel animation plays and lands on the winning segment.')
add_step(7, 'The visitor sees their prize and coupon code.')

add_heading('Setting Up Prize Segments', 2)
add_para(
    'When you create a Spin-to-Win form, an extra Spin Config tab appears in the form builder between '
    'the Fields and Design steps. Here you configure up to 8 prize segments that appear on the wheel.'
)

add_step(1, 'Go to the Spin Config tab in the form builder.')
add_step(2, 'Click + Add Segment to add a new prize segment.')
add_step(3, 'For each segment, configure:')
add_bullet('A hex colour for the wheel segment.', 'Colour:')
add_bullet('The text shown on the wheel (e.g., "10% Off", "Free Shipping", "Try Again").', 'Label:')
add_bullet('A number that determines how likely this segment is to win (higher = more likely). These are relative weights, not percentages.', 'Probability Weight:')
add_bullet('Choose % Off (percentage discount), $ Off (fixed amount discount), or Free Shipping.', 'Discount Type:')
add_bullet('The discount amount (e.g., 10 for 10% off or 5 for $5 off). Not applicable for Free Shipping.', 'Discount Value:')

add_step(4, 'Repeat for each prize segment (up to 8 maximum).')
add_step(5, 'Click Save Form when done.')

add_tip(
    'Include at least one "no prize" segment (e.g., "Try Again" with a discount value of 0) to create '
    'suspense and control how often visitors win. Set it with a higher probability weight.'
)

add_heading('How Probability Weights Work', 2)
add_para(
    'The probability of landing on a segment is its weight divided by the total weight of all segments. '
    'For example, if you have three segments:'
)
add_bullet('"10% Off" with weight 20')
add_bullet('"Free Shipping" with weight 10')
add_bullet('"Try Again" with weight 70')
add_para(
    'The total weight is 100. The chances are: 10% Off = 20%, Free Shipping = 10%, Try Again = 70%. '
    'You can use any numbers \u2014 the plugin calculates the ratios automatically.'
)

add_heading('Automatic Coupon Generation', 2)
add_para(
    'When a visitor wins a prize with a discount value (or free shipping), the plugin automatically '
    'creates a WooCommerce coupon with the following settings:'
)
add_bullet('Code format: AMS-SPIN-XXXXXXXX (8 random characters).')
add_bullet('Single-use: Each coupon can only be used once.')
add_bullet('Email-restricted: The coupon is locked to the subscriber\'s email address.')
add_bullet('Expiry: 30 days by default (configurable per segment via the expiry_days field).')
add_bullet('Minimum spend: Configurable per segment (optional).')
add_bullet('Discount type: Percentage or fixed cart discount, matching the segment configuration.')

add_important(
    'Coupon generation requires WooCommerce to be present. If you are running in standalone mode '
    '(without WooCommerce), the spin-to-win form will still work but coupon codes will not be generated. '
    'The subscriber will still be captured.'
)

add_heading('The Visitor Experience', 2)
add_step(1, 'A modal pop-up appears with the spin wheel on one side and a form on the other.')
add_step(2, 'The visitor enters their email (and any other fields you configured).')
add_step(3, 'They click the submit button (e.g., "Spin the Wheel!").')
add_step(4, 'The wheel spins with an animation and lands on a segment.')
add_step(5, 'If they won a prize, they see the prize label and coupon code prominently displayed.')
add_step(6, 'They can copy the coupon code and use it at checkout.')

add_tip(
    'For the best conversion rates, keep the form fields minimal \u2014 just Email and optionally First Name. '
    'The goal is to make it quick and fun.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 13: SMS MARKETING
# ============================================================
add_heading('Chapter 13: SMS Marketing', 1)

add_para(
    'Apotheca Marketing Suite includes built-in SMS marketing powered by Twilio. You can send SMS and MMS '
    'campaigns to your subscribers, use personalisation tokens, track delivery status, and manage TCPA '
    'compliance with automatic STOP/UNSTOP/HELP keyword handling.'
)

add_heading('Setting Up Twilio Integration', 2)
add_para(
    'Before you can send SMS messages, you need a Twilio account and a phone number. Here is how to '
    'set everything up:'
)

add_heading('Getting Your Twilio Credentials', 3)
add_step(1, 'Go to twilio.com and create an account (or log into your existing account).')
add_step(2, 'From the Twilio Console dashboard, locate your Account SID and Auth Token. These are displayed prominently on the main console page.')
add_step(3, 'Purchase a phone number from the Twilio Phone Numbers section. This will be your "From" number. Choose a number that supports SMS (and MMS if you plan to send images).')
add_step(4, 'Alternatively, if you have a Twilio Messaging Service, you can use the Messaging Service SID (starts with "MG") instead of a phone number.')

add_heading('Entering Credentials in the Plugin', 3)
add_step(1, 'In the WordPress admin, go to Apotheca Marketing > SMS.')
add_step(2, 'Click the SMS Settings tab at the top.')
add_step(3, 'Enter your Account SID in the Account SID field.')
add_step(4, 'Enter your Auth Token in the Auth Token field.')
add_step(5, 'Enter your Twilio phone number (e.g., +12125551234) in the From Number / Messaging Service SID field. If you are using a Messaging Service, enter its SID (e.g., MGxxxxxxxx) instead.')
add_step(6, 'Optionally customise the HELP Reply Text \u2014 this is the automated response sent when a subscriber texts HELP to your number.')
add_step(7, 'Click Save Credentials.')

add_para(
    'A green status message will confirm "SMS provider is configured and ready." If there is a problem, '
    'you will see a yellow warning.'
)

add_heading('How Credentials Are Encrypted', 3)
add_para(
    'Your Twilio credentials are encrypted using AES-256-CBC before being stored in the database. The '
    'encryption key is derived from your WordPress AUTH_KEY constant (defined in wp-config.php). This means:'
)
add_bullet('Credentials are never stored in plain text.')
add_bullet('Even if someone accesses the database directly, they cannot read the credentials without your AUTH_KEY.')
add_bullet('Each credential (Account SID, Auth Token, From Number) is encrypted independently with a unique random IV (initialisation vector).')

add_important(
    'Make sure your wp-config.php has a strong, unique AUTH_KEY set. WordPress generates this by default '
    'during installation.'
)

add_heading('Configuring Twilio Webhooks', 3)
add_para(
    'The SMS Settings page displays two webhook URLs that you need to configure in your Twilio account:'
)
add_bullet('The URL that handles inbound messages (STOP, UNSTOP, HELP keywords). Copy this URL and paste it into your Twilio phone number\'s Messaging > "A message comes in" webhook field.', 'Inbound Webhook:')
add_bullet('The URL that receives delivery status updates from Twilio. The plugin sets this automatically on each outbound message, so you typically do not need to configure it manually.', 'Status Callback:')

add_heading('Sending a Test SMS', 3)
add_step(1, 'In the SMS Settings tab, scroll down to the Test Send section.')
add_step(2, 'Enter a phone number in the field (use international format, e.g., +12125551234).')
add_step(3, 'Click Send Test SMS.')
add_step(4, 'You should receive a test message on that phone within seconds. A confirmation message will appear on screen.')

add_heading('Creating an SMS Campaign', 2)
add_step(1, 'Go to Apotheca Marketing > SMS.')
add_step(2, 'Click the Campaigns tab (selected by default).')
add_step(3, 'Click + New Campaign.')
add_step(4, 'Enter a Campaign Name (e.g., "Flash Sale SMS").')
add_step(5, 'Optionally, enter a Segment ID to target a specific subscriber segment. Leave blank to send to all SMS-opted-in subscribers.')
add_step(6, 'Compose your message in the Message Body textarea. You can click the personalisation token buttons above the textarea to insert tokens.')
add_step(7, 'Optionally, enter an MMS Image URL to send a picture message (MMS). The URL must point to a publicly accessible image file (JPEG, PNG, or GIF).')
add_step(8, 'Click Save to save the campaign as a draft.')
add_step(9, 'To send, go back to the campaign list and click Send Now next to your campaign. You will be asked to confirm.')

add_heading('Personalisation Tokens', 2)
add_para('The following personalisation tokens are available in SMS messages:')
add_bullet('The subscriber\'s first name (defaults to "there" if not set).', '{{first_name}}:')
add_bullet('The subscriber\'s last name.', '{{last_name}}:')
add_bullet('The subscriber\'s email address.', '{{email}}:')
add_bullet('The subscriber\'s phone number.', '{{phone}}:')
add_bullet('Your store\'s name (from WordPress Settings > General).', '{{shop_name}}:')
add_bullet('Your store\'s URL.', '{{shop_url}}:')
add_bullet('A link to unsubscribe from SMS.', '{{unsubscribe_url}}:')
add_bullet('The customer\'s order number (for order-triggered flows).', '{{order_number}}:')
add_bullet('The order total (for order-triggered flows).', '{{order_total}}:')
add_bullet('The product name (for product-specific flows).', '{{product_name}}:')
add_bullet('A link to the customer\'s cart (for abandoned cart recovery).', '{{cart_url}}:')
add_bullet('A coupon code (when used in flows with coupon generation).', '{{coupon_code}}:')

add_heading('Character Counter and SMS Segments', 2)
add_para(
    'The campaign editor includes a live character counter. A standard SMS message is 160 characters. '
    'If your message contains personalisation tokens, the limit shown is 153 characters (to account for '
    'variable-length token replacements). Messages longer than the limit will be split into multiple SMS '
    'segments, which may increase your Twilio costs.'
)
add_tip(
    'Keep messages concise. The character counter shows the segment count (e.g., "165/160 chars (2 SMS segments)"). '
    'Try to fit within a single segment when possible.'
)

add_heading('Message Preview', 2)
add_para(
    'Below the message body, a live Preview shows what the message will look like with sample data '
    'replacing the tokens (e.g., {{first_name}} becomes "Jane", {{order_total}} becomes "$49.99"). '
    'The auto-appended "Reply STOP to unsubscribe." footer is also shown.'
)

add_heading('TCPA Compliance', 2)
add_para(
    'The plugin includes built-in TCPA (Telephone Consumer Protection Act) compliance features:'
)

add_heading('Automatic Keyword Handling', 3)
add_bullet('When a subscriber texts STOP to your number, they are automatically opted out of SMS. Their sms_opt_in flag is set to 0 and they will not receive any further SMS messages.', 'STOP:')
add_bullet('When a previously opted-out subscriber texts UNSTOP (or START), they are re-opted in to SMS. Their sms_opt_in flag is set back to 1.', 'UNSTOP / START:')
add_bullet('When a subscriber texts HELP, they receive your configured HELP Reply Text (e.g., "Reply STOP to opt out. Reply HELP for help. Msg&data rates may apply.").', 'HELP:')

add_heading('Opt-In Requirement', 3)
add_para(
    'The plugin checks the subscriber\'s sms_opt_in flag before every send. If the flag is not set to 1, '
    'the message is skipped and the send record is marked as "skipped". This prevents messages from being '
    'sent to subscribers who have not consented to SMS.'
)

add_heading('Automatic Footer', 3)
add_para(
    '"Reply STOP to unsubscribe." is automatically appended to every outbound SMS message.'
)

add_heading('SMS Consent Management', 2)
add_para(
    'SMS consent is managed separately from email consent. A subscriber can be opted in to email but opted '
    'out of SMS, or vice versa. The sms_opt_in flag on the subscriber record controls SMS eligibility.'
)
add_para('Subscribers can opt in to SMS through:')
add_bullet('Including a Phone field on an opt-in form along with appropriate consent text.')
add_bullet('Texting a keyword to your Twilio number (if you set up an opt-in keyword).')
add_bullet('Manual opt-in by an admin in the Subscribers section.')

add_heading('MMS Support', 2)
add_para(
    'To send an MMS (picture message), enter a publicly accessible image URL in the MMS Image URL field '
    'when creating a campaign. Supported formats include JPEG, PNG, and GIF. The image URL is passed to '
    'Twilio\'s MediaUrl parameter.'
)
add_important(
    'MMS is only available in the US and Canada. MMS messages cost more than standard SMS. '
    'Check your Twilio pricing for details.'
)

add_heading('Delivery Tracking and Retry Logic', 2)
add_para(
    'Every SMS send is tracked with a delivery status. The plugin registers a Status Callback URL with '
    'Twilio so that delivery updates (delivered, undelivered, failed) are received automatically.'
)
add_para('Send statuses include:')
add_bullet('The message is waiting to be sent.', 'queued:')
add_bullet('The message was accepted by Twilio.', 'sent:')
add_bullet('A retry has been scheduled after an initial failure.', 'retry_queued:')
add_bullet('The send was skipped (subscriber not opted in).', 'skipped:')
add_bullet('The message failed after the retry attempt.', 'permanently_failed:')

add_para(
    'If a send fails, the plugin schedules one retry after 30 minutes. If the retry also fails, the '
    'status is set to permanently_failed.'
)

add_heading('Viewing Send History', 2)
add_para(
    'All SMS sends are recorded in the ams_sends table with channel set to "sms". You can view send '
    'history and delivery statuses from the campaign list or in the Analytics dashboard.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 14: VISUAL EMAIL EDITOR
# ============================================================
add_heading('Chapter 14: Visual Email Editor', 1)

add_para(
    'The Visual Email Editor is a drag-and-drop block editor for creating email templates. You build emails '
    'by adding, arranging, and customising blocks. The editor produces fully responsive, email-client-compatible '
    'HTML with CSS inlining and Outlook-specific VML fallbacks.'
)

add_heading('Navigating to the Email Editor', 2)
add_step(1, 'In the WordPress admin, go to Apotheca Marketing > Email Editor.')
add_step(2, 'You will see the Email Templates list showing all saved templates.')
add_step(3, 'Click + New Template to create a new email, or click Edit on an existing template.')
add_step(4, 'You can also Duplicate a template to create a copy, or Delete templates you no longer need.')

add_heading('The Editor Interface', 2)
add_para('The editor has three main panels:')
add_bullet('Contains the Blocks tab (block palette) and the Style tab (global style settings).', 'Left Panel:')
add_bullet('The email canvas where you arrange blocks. You can drag blocks from the palette and drop them here, or click a block in the palette to add it to the bottom. Click any block on the canvas to select it. Use the up/down arrows to reorder blocks.', 'Centre Panel (Canvas):')
add_bullet('Shows the settings for the currently selected block. Click a block on the canvas to see its settings here.', 'Right Panel (Block Settings):')

add_para('The toolbar at the top provides:')
add_bullet('A back arrow to return to the template list.')
add_bullet('An editable template name field.')
add_bullet('A Preview button that renders the full email HTML and shows it in an iframe.')
add_bullet('A Save button to save your work.')

add_heading('The 12 Block Types', 2)

add_heading('Header Block', 3)
add_para('Displays your brand logo at the top of the email.')
add_bullet('Logo URL \u2014 The URL of your logo image.')
add_bullet('Logo Width \u2014 The display width in pixels.')
add_bullet('Alignment \u2014 Left, Centre, or Right.')

add_heading('Text Block', 3)
add_para('A rich text area for paragraphs, headings, lists, and inline formatting.')
add_bullet('Content \u2014 Enter HTML content (paragraphs, bold, italic, links, etc.).')
add_bullet('Font Size \u2014 The text size (e.g., 16px).')
add_bullet('Alignment \u2014 Left, Centre, or Right.')

add_heading('Image Block', 3)
add_para('Displays an image in the email.')
add_bullet('Image URL \u2014 The URL of the image.')
add_bullet('Alt Text \u2014 Alternative text for accessibility and when images are blocked.')
add_bullet('Width \u2014 The display width (e.g., 100%, 300px).')
add_bullet('Link URL \u2014 An optional clickable link when the image is clicked.')
add_bullet('Alignment \u2014 Left, Centre, or Right.')

add_heading('Button Block', 3)
add_para('A call-to-action button with full customisation.')
add_bullet('Button Text \u2014 The text on the button.')
add_bullet('URL \u2014 The link destination.')
add_bullet('Background Color and Text Color.')
add_bullet('Border Radius \u2014 Roundness of the button corners.')
add_bullet('Button Padding \u2014 Internal spacing (e.g., 12px 30px).')
add_bullet('Alignment \u2014 Left, Centre, or Right.')
add_para(
    'The button block includes Outlook-specific VML (Vector Markup Language) code so that buttons '
    'render correctly in Microsoft Outlook, which does not support CSS border-radius or padding on links.'
)

add_heading('Divider Block', 3)
add_para('A horizontal line to visually separate sections.')
add_bullet('Color \u2014 The line colour.')
add_bullet('Width \u2014 The line width (e.g., 100%, 80%).')
add_bullet('Height \u2014 The line thickness (e.g., 1px, 2px).')

add_heading('Spacer Block', 3)
add_para('Adds vertical whitespace between blocks.')
add_bullet('Height \u2014 The spacer height (e.g., 20px, 40px).')

add_heading('Columns Block', 3)
add_para(
    'Creates a multi-column layout for side-by-side content. Columns stack vertically on mobile.'
)
add_bullet('50/50 \u2014 Two equal-width columns.', 'Layout Options:')
add_bullet('60/40 \u2014 Wider left column.')
add_bullet('40/60 \u2014 Wider right column.')
add_bullet('33/33/33 \u2014 Three equal-width columns.')
add_para(
    'Each column can contain other blocks. The Columns block uses MSO conditional comments to render '
    'correctly in Outlook, which does not support CSS inline-block.'
)

add_heading('Product Block', 3)
add_para(
    'Displays a WooCommerce product card with image, name, price, and a link to the product page.'
)
add_bullet('Product ID \u2014 The WooCommerce product ID to display.')
add_important('The Product block requires WooCommerce to be present. In standalone mode, this block will not render.')

add_heading('Social Block', 3)
add_para('Displays social media icon links in a row.')
add_bullet('Alignment \u2014 Left, Centre, or Right.')
add_bullet('Icon Size \u2014 The size of each icon in pixels.')
add_bullet('Social Links \u2014 A list of platform/URL pairs. Supported platforms: Facebook, Twitter, Instagram, LinkedIn, YouTube, TikTok, Pinterest.')

add_heading('Reviews Block', 3)
add_para('Displays product reviews in the email. Three modes are available:')
add_bullet('Shows existing customer reviews with star ratings. You specify a product ID and maximum number of reviews.', 'Social Proof:')
add_bullet('Asks the recipient to leave a review with a customisable CTA button.', 'Review Request:')
add_bullet('Presents "Great!" and "Could be better" buttons that route to different destinations (covered in Chapter 15).', 'Review Gating:')

add_heading('Footer Block', 3)
add_para(
    'A text area for footer content like copyright notices, unsubscribe links, and physical address.'
)
add_bullet('Content \u2014 HTML content. The default includes {{shop_name}} and {{unsubscribe_url}} tokens.')

add_heading('HTML Block', 3)
add_para(
    'A raw HTML block for advanced users. You can paste any HTML code directly. Use this for custom '
    'content that the other block types cannot produce.'
)

add_heading('Personalisation Tokens in Emails', 2)
add_para(
    'You can use personalisation tokens anywhere in text content, button URLs, and other text fields. '
    'Common tokens include:'
)
add_bullet('{{first_name}}, {{last_name}}, {{email}}')
add_bullet('{{shop_name}}, {{shop_url}}')
add_bullet('{{unsubscribe_url}} \u2014 One-click unsubscribe link')
add_bullet('{{order_number}}, {{order_total}}, {{order_status}}')
add_bullet('{{cart_url}}, {{coupon_code}}, {{product_name}}')

add_heading('Google Fonts Support', 2)
add_para(
    'The email editor supports Google Fonts with a three-tier fallback system for maximum compatibility:'
)
add_bullet('Webmail clients (Gmail, Yahoo) that support @import will load the Google Font (Montserrat by default) via a CSS @import rule in the email\'s <head>.', 'Tier 1 \u2014 Google Fonts:')
add_bullet('Microsoft Outlook ignores @import but supports Century Gothic, which is visually similar to Montserrat. The email includes an Outlook-specific MSO conditional style that forces Century Gothic.', 'Tier 2 \u2014 Outlook Fallback:')
add_bullet('All other clients fall back to the system sans-serif font.', 'Tier 3 \u2014 System Fallback:')

add_para('In the Design step of the form builder (and global style settings in the email editor), you can choose from:')
add_bullet('Montserrat, Open Sans, Roboto, Lato, Poppins, Playfair Display, Merriweather, Raleway, Nunito, Inter, or System Default.')

add_heading('Live Preview', 2)
add_para(
    'Click the Preview button in the toolbar to see a rendered version of your email. The plugin sends '
    'the blocks and global style to the server, which renders the full HTML (including CSS inlining) and '
    'displays it in an iframe. This gives you an accurate representation of how the email will look in '
    'most email clients.'
)

add_heading('How CSS Inlining Works', 2)
add_para(
    'Most email clients (especially Gmail) strip <style> tags from emails. To ensure your styles are '
    'applied, the plugin converts CSS rules into inline style attributes on each HTML element. Here is '
    'how it works:'
)
add_step(1, 'The renderer outputs the email with CSS in <style> tags.')
add_step(2, 'The CSS Inliner extracts all <style> blocks from the HTML.')
add_step(3, '@import rules (for Google Fonts) and @media rules (for mobile responsiveness) are preserved and placed back in the <head>.')
add_step(4, 'All other CSS rules are parsed and matched to HTML elements using a CSS-to-XPath converter.')
add_step(5, 'The matched styles are applied as inline style attributes on each element.')
add_step(6, 'The final HTML has both inline styles (for compatibility) and @media rules (for mobile responsiveness).')

add_heading('Outlook Compatibility', 2)
add_para(
    'Microsoft Outlook uses the Word rendering engine, which has limited CSS support. The email renderer '
    'includes several Outlook-specific features:'
)
add_bullet('VML (Vector Markup Language) roundrect elements for buttons, ensuring they display correctly with background colours and rounded corners.')
add_bullet('MSO conditional comments (<!--[if mso]>) for table-based column layouts.')
add_bullet('Century Gothic font fallback in MSO-specific <style> blocks.')
add_bullet('PixelsPerInch setting to ensure consistent rendering at 96 DPI.')

add_heading('Mobile Responsiveness', 2)
add_para(
    'The email renderer includes @media queries that stack columns vertically on screens narrower than '
    '620px. The .ams-col class switches from inline-block to block display and from a percentage width '
    'to 100% width. This ensures your multi-column layouts read well on mobile devices.'
)

add_heading('Saving and Using Templates', 2)
add_para(
    'Templates are saved independently and can be used across campaigns and flows. When you create a '
    'campaign or a flow step that sends an email, you select a template from your saved templates.'
)
add_step(1, 'In the email editor, build and customise your template.')
add_step(2, 'Give it a descriptive name (e.g., "Welcome Email", "Abandoned Cart Recovery").')
add_step(3, 'Click Save.')
add_step(4, 'When creating a campaign or flow email step, select this template by name.')

add_tip(
    'Use the Duplicate button to create variations of a template (e.g., for A/B testing subject lines) '
    'without starting from scratch.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 15: REVIEWS INTEGRATION
# ============================================================
add_heading('Chapter 15: Reviews Integration', 1)

add_para(
    'Apotheca Marketing Suite integrates with WooCommerce native reviews and the Judge.me third-party '
    'review platform. Reviews are cached locally and can be displayed in your marketing emails as social '
    'proof. The plugin also includes a review gating feature that routes positive ratings to public review '
    'pages and negative ratings to private feedback forms.'
)

add_heading('Setting Up WooCommerce Reviews Import', 2)
add_para(
    'If WooCommerce is installed on the same site as Apotheca Marketing Suite (co-located mode), '
    'WooCommerce reviews are imported automatically. No additional configuration is required.'
)
add_step(1, 'Go to Apotheca Marketing > Reviews.')
add_step(2, 'The Review Cache card shows the number of cached reviews and their sources.')
add_step(3, 'Set the Minimum Rating to Cache setting (default is 4 Stars & Above). This controls which reviews are imported \u2014 lower-rated reviews are excluded.')
add_step(4, 'Click Save Settings.')
add_step(5, 'Click Refresh Now to trigger an immediate import.')

add_para(
    'The importer pulls approved reviews from WooCommerce in batches of 200, filters them by the '
    'minimum rating, and stores them in the ams_reviews_cache table. Duplicate reviews are detected '
    'and skipped automatically.'
)

add_tip(
    'In standalone mode (no WooCommerce), the WooCommerce importer is automatically skipped. '
    'Reviews can still come from Judge.me if configured.'
)

add_heading('Setting Up Judge.me Integration (Optional)', 2)
add_para(
    'If you use Judge.me for product reviews, you can import those reviews into the cache as well.'
)

add_step(1, 'Make sure the Judge.me plugin is installed and active on your site.')
add_step(2, 'Go to Apotheca Marketing > Reviews.')
add_step(3, 'In the Judge.me Integration card, the status will show "Plugin Detected \u2014 API Key Required".')
add_step(4, 'Enter your Judge.me API key in the API Key field. You can find this in your Judge.me dashboard under Settings > API.')
add_step(5, 'Click Test Connection to verify the API key works. You should see "Connection successful."')
add_step(6, 'Click Save Settings.')

add_para(
    'Once configured, Judge.me reviews are imported alongside WooCommerce reviews during the nightly '
    'cache refresh. The importer pulls reviews via the Judge.me REST API, filtered by your minimum '
    'rating setting.'
)

add_heading('How Review Gating Works', 2)
add_para(
    'Review gating is a technique where you ask customers how satisfied they are before directing them '
    'to leave a public review. Customers who had a positive experience (4-5 stars) are directed to your '
    'public review page. Customers who had a negative experience (1-3 stars) are directed to a private '
    'feedback form where you can address their concerns personally.'
)

add_heading('The Review Gate Flow', 3)
add_step(1, 'After a customer completes an order, a post-purchase email is sent (via an automated flow).')
add_step(2, 'The email contains a review gating block with two buttons: a positive button (e.g., "Great!") and a negative button (e.g., "Could be better"). Alternatively, it shows 1-5 star rating links.')
add_step(3, 'When the customer clicks a button, they are directed to your site\'s /ams-review-gate/ endpoint with their token, order ID, and selected rating.')
add_step(4, 'The plugin validates the token, verifies the order belongs to the subscriber, and checks that the gate link has not expired or been used before.')
add_step(5, 'For ratings of 4-5 stars: The customer is redirected to the product review page on your store (e.g., the product page\'s Reviews tab).')
add_step(6, 'For ratings of 1-3 stars: The customer is redirected to a private feedback page that you configure in settings. If no feedback page is set, they see a simple "Thank You for Your Feedback" page.')

add_heading('Configuring Review Gating', 3)
add_step(1, 'Go to Apotheca Marketing > Reviews.')
add_step(2, 'In Review Settings, set the Private Feedback Page by selecting a published WordPress page from the dropdown. This is where 1-3 star customers will be directed.')
add_step(3, 'Set the Review Gate Link Expiry (hours) \u2014 the default is 72 hours. After this period, the gate link in the email will no longer work.')
add_step(4, 'Click Save Settings.')
add_step(5, 'In your post-purchase flow, add a Reviews block to your email template with the mode set to "Review Gating".')

add_important(
    'Review gate links are single-use. Once a customer clicks a link, it cannot be used again. '
    'This prevents manipulation and ensures accurate tracking.'
)

add_heading('The Reviews Block in the Email Editor', 2)
add_para(
    'The Reviews block (available in the Visual Email Editor \u2014 see Chapter 14) has three modes:'
)

add_heading('Social Proof Mode', 3)
add_para(
    'Displays existing customer reviews in the email body. You configure a product ID and the maximum '
    'number of reviews to show. The block renders star ratings, review text (trimmed to 30 words), '
    'reviewer name, and a verified purchase badge if applicable.'
)

add_heading('Review Request Mode', 3)
add_para(
    'Displays a heading, body text, and a call-to-action button asking the customer to leave a review. '
    'You customise the heading (e.g., "How was your experience?"), body text, CTA button text '
    '(e.g., "Leave a Review"), and the CTA URL (typically {{review_url}} which resolves to the product page).'
)

add_heading('Review Gating Mode', 3)
add_para(
    'Displays a heading and two buttons \u2014 a positive button (green, e.g., "Great!") and a negative '
    'button (grey, e.g., "Could be better"). Each button links to a different URL. The positive URL '
    'typically points to the public review page, and the negative URL points to the private feedback form.'
)

add_heading('How Contextual Review Selection Works', 2)
add_para(
    'When the Reviews block renders in an email, the plugin intelligently selects the most relevant '
    'reviews based on the flow context. There are four selection modes:'
)

add_heading('Auto-Contextual (Default)', 3)
add_para('The plugin automatically picks reviews based on the flow trigger type:')
add_bullet('Fetches reviews for the products in the subscriber\'s most recent abandoned cart.', 'Abandoned Cart Emails:')
add_bullet('Fetches reviews for products in the subscriber\'s most frequently purchased category.', 'Win-Back Emails:')
add_bullet('Falls back to sitewide verified 5-star reviews.', 'All Other Emails:')

add_heading('Specific Product', 3)
add_para('Fetches reviews for a specific product ID that you set in the block settings. Useful for product launch or upsell emails.')

add_heading('Top Rated Sitewide', 3)
add_para('Fetches the highest-rated reviews across all products. Good for general social proof in welcome emails or newsletters.')

add_heading('Most Recent Sitewide', 3)
add_para('Fetches the most recent reviews (rated 4 stars or above) across all products. Good for showing fresh, active reviews.')

add_heading('Review Card Design', 2)
add_para('Each review card in the email includes:')
add_bullet('Star Rating \u2014 Gold star icons (filled for the rating, empty for the remainder).')
add_bullet('Review Title \u2014 If the review has a title (more common with Judge.me reviews).')
add_bullet('Review Body \u2014 The review text, trimmed to 30 words for email brevity.')
add_bullet('Reviewer Name \u2014 Displayed as "\u2014 Jane Doe".')
add_bullet('Verified Purchase Badge \u2014 A green checkmark with "Verified" if the reviewer is a verified buyer.')
add_bullet('Product Name \u2014 Shown above the stars, linked to the product page if a URL is available.')
add_bullet('Max Reviews \u2014 Configurable in the block settings (default: 3).')

add_heading('Nightly Cache Refresh', 2)
add_para(
    'The review cache is refreshed automatically every night at 3:00 AM via Action Scheduler. '
    'The nightly job performs the following steps:'
)
add_step(1, 'Purges stale cache entries older than 48 hours.')
add_step(2, 'Imports WooCommerce native reviews (if WooCommerce is present) in batches of 200.')
add_step(3, 'Imports Judge.me reviews (if the Judge.me plugin is active and an API key is configured).')
add_step(4, 'Updates the "last refresh" timestamp.')

add_para(
    'You can also trigger a manual refresh at any time by clicking Refresh Now on the Reviews settings page.'
)

add_heading('Review Settings Configuration Summary', 2)
add_para('Here is a summary of all configurable review settings:')

# Settings summary table
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
headers = ['Setting', 'Description']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows_data = [
    ('Minimum Rating to Cache', 'Controls which reviews are imported. Options: 3 Stars & Above, 4 Stars & Above (default), or 5 Stars Only.'),
    ('Private Feedback Page', 'The WordPress page where customers with 1-3 star ratings are redirected. Select "None" to show a built-in thank-you page.'),
    ('Review Gate Link Expiry', 'How many hours a review gate link remains valid after the order is completed. Default: 72 hours. Maximum: 720 hours (30 days).'),
    ('Judge.me API Key', 'Your Judge.me API key for importing reviews from the Judge.me platform. Encrypted with AES-256-CBC before storage.'),
]
for i, (setting, desc) in enumerate(rows_data):
    table.rows[i + 1].cells[0].text = setting
    table.rows[i + 1].cells[1].text = desc

add_para('')  # spacing after table

add_tip(
    'Start with "4 Stars & Above" as the minimum rating. This ensures only positive reviews appear '
    'in your marketing emails, which maximises social proof effectiveness.'
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Apotheca_User_Manual_Part3.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
